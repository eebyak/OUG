#!/usr/bin/env python3
"""Build readable Word and Markdown module handbooks from OOAPI-style JSON files.

Example
-------
python build_module_handbook.py \
    --dir ./dir \
    --pattern 'T*INF*v6.json' \
    --codebook ./competence_codebook_v11_ooapi.json \
    --output ./OUG_Module_Handbook.docx \
    --markdown-output ./OUG_Module_Handbook.md \
    --language en-GB \
    --institution 'Open University of Germany (OUG)' \
    --programme 'Bachelor of Computer Science (B.Sc.)'

Both outputs deliberately render only human-facing content. Internal UUIDs,
extension tags, processing metadata, legacy mappings, API identifiers and similar
machine-oriented fields are not copied into the handbook.

Dependency: python-docx >= 1.1
Install with: python -m pip install python-docx
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - helpful CLI error
    raise SystemExit(
        "python-docx is required. Install it with: python -m pip install python-docx"
    ) from exc


# ---------------------------------------------------------------------------
# Visual system
# ---------------------------------------------------------------------------
TEAL = "2F5D62"
DARK_TEAL = "173F43"
PALE_TEAL = "EAF2F2"
PALE_GREY = "F4F6F7"
MID_GREY = "667276"
DARK_GREY = "263238"
WHITE = "FFFFFF"
ACCENT = "D8A43B"
FONT = "Arial"


LANGUAGE_LABELS = {
    "eng": {"en": "English", "de": "Englisch"},
    "ger": {"en": "German", "de": "Deutsch"},
    "deu": {"en": "German", "de": "Deutsch"},
    "fra": {"en": "French", "de": "Französisch"},
    "spa": {"en": "Spanish", "de": "Spanisch"},
}

DELIVERY_LABELS = {
    "online": {"en": "Online", "de": "Online"},
    "hybrid": {"en": "Hybrid", "de": "Hybrid"},
    "onsite": {"en": "On campus", "de": "Präsenz"},
    "on-campus": {"en": "On campus", "de": "Präsenz"},
    "blended": {"en": "Blended", "de": "Blended"},
}

COMPLEXITY_LABELS = {
    "remember": {"en": "Foundation", "de": "Grundlage"},
    "understand": {"en": "Understanding & application", "de": "Verstehen & Anwenden"},
    "apply": {"en": "Application", "de": "Anwendung"},
    "analyze": {"en": "Advanced analysis", "de": "Vertiefte Analyse"},
    "analyse": {"en": "Advanced analysis", "de": "Vertiefte Analyse"},
    "evaluate": {"en": "Evaluation & creation", "de": "Bewerten & Gestalten"},
    "create": {"en": "Evaluation & creation", "de": "Bewerten & Gestalten"},
}

UI = {
    "en": {
        "assembled": "Assembled",
        "source_note": "Human-readable edition assembled from machine-readable module data.",
        "collection_overview": "Collection overview",
        "modules": "Modules",
        "total_ects": "Total ECTS",
        "streams": "Streams",
        "academic_years": "Academic years",
        "module_index": "Module index",
        "code": "Code",
        "module": "Module",
        "stream": "Stream",
        "year": "Year",
        "ects": "ECTS",
        "at_a_glance": "At a glance",
        "workload": "Total workload",
        "contact": "Contact time",
        "self_study": "Self-study",
        "level": "Level",
        "duration": "Duration",
        "delivery": "Delivery",
        "languages": "Languages",
        "assessment": "Assessment",
        "overview": "Module overview",
        "learning_outcomes": "Learning outcomes",
        "competence_area": "Competence area",
        "resources": "Recommended literature and resources",
        "not_specified": "Not specified",
        "hours": "hours",
        "month": "month",
        "months": "months",
        "year_singular": "year",
        "years": "years",
        "source": "Source file",
        "generated": "Generated from structured module data; internal identifiers and technical metadata have been omitted.",
    },
    "de": {
        "assembled": "Zusammengestellt am",
        "source_note": "Lesefassung aus maschinenlesbaren Moduldaten.",
        "collection_overview": "Überblick",
        "modules": "Module",
        "total_ects": "ECTS gesamt",
        "streams": "Studienströme",
        "academic_years": "Studienjahre",
        "module_index": "Modulübersicht",
        "code": "Code",
        "module": "Modul",
        "stream": "Studienstrom",
        "year": "Jahr",
        "ects": "ECTS",
        "at_a_glance": "Auf einen Blick",
        "workload": "Gesamtaufwand",
        "contact": "Kontaktzeit",
        "self_study": "Selbststudium",
        "level": "Niveau",
        "duration": "Dauer",
        "delivery": "Durchführung",
        "languages": "Sprachen",
        "assessment": "Prüfungsform",
        "overview": "Modulbeschreibung",
        "learning_outcomes": "Lernergebnisse",
        "competence_area": "Kompetenzbereich",
        "resources": "Empfohlene Literatur und Ressourcen",
        "not_specified": "Nicht angegeben",
        "hours": "Stunden",
        "month": "Monat",
        "months": "Monate",
        "year_singular": "Jahr",
        "years": "Jahre",
        "source": "Quelldatei",
        "generated": "Aus strukturierten Moduldaten erzeugt; interne Kennungen und technische Metadaten wurden weggelassen.",
    },
}


# ---------------------------------------------------------------------------
# Data helpers
# ---------------------------------------------------------------------------
def base_language(language: str) -> str:
    return "de" if language.lower().startswith("de") else "en"


def ui(language: str, key: str) -> str:
    return UI[base_language(language)][key]


def load_json(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError(f"Expected a JSON object in {path}")
    return data


def localized_value(
    value: Any,
    language: str,
    *,
    fallback_languages: Sequence[str] = ("en-GB", "de-DE", "en", "de"),
) -> str:
    """Return a human string from OOAPI multilingual values or plain strings."""
    if value is None:
        return ""
    if isinstance(value, str):
        return clean_text(value)
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, Mapping):
        if "value" in value:
            return clean_text(value.get("value", ""))
        return ""
    if not isinstance(value, Sequence):
        return clean_text(str(value))

    items = [item for item in value if isinstance(item, Mapping)]
    if not items:
        return ""

    requested = language.lower()
    requested_base = requested.split("-")[0]

    for item in items:
        item_lang = str(item.get("language", "")).lower()
        if item_lang == requested:
            return clean_text(item.get("value", ""))
    for item in items:
        item_lang = str(item.get("language", "")).lower()
        if item_lang.split("-")[0] == requested_base:
            return clean_text(item.get("value", ""))
    for fallback in fallback_languages:
        fallback = fallback.lower()
        for item in items:
            item_lang = str(item.get("language", "")).lower()
            if item_lang == fallback or item_lang.split("-")[0] == fallback.split("-")[0]:
                return clean_text(item.get("value", ""))
    return clean_text(items[0].get("value", ""))


def clean_text(value: Any) -> str:
    """Normalise whitespace without changing substantive wording."""
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def natural_sort_key(text: str) -> list[Any]:
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", text)]


def number(value: Any) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def pretty_number(value: float | int | None) -> str:
    if value is None:
        return ""
    if float(value).is_integer():
        return str(int(value))
    return f"{value:g}"


def study_load_map(module: Mapping[str, Any]) -> dict[str, float]:
    result: dict[str, float] = {}
    for entry in module.get("studyLoad", []) or []:
        if not isinstance(entry, Mapping):
            continue
        unit = str(entry.get("studyLoadUnit", "")).lower()
        val = number(entry.get("value"))
        if unit and val is not None:
            result[unit] = val

    ext_workload = ((module.get("ext") or {}).get("x-dhbw-workload") or {})
    if isinstance(ext_workload, Mapping):
        aliases = {"total": "sbu", "contact": "contact_time", "selfStudy": "self_study"}
        for key, unit in aliases.items():
            val = number(ext_workload.get(key))
            if val is not None and unit not in result:
                result[unit] = val
    if "self_study" not in result and "sbu" in result and "contact_time" in result:
        result["self_study"] = max(0, result["sbu"] - result["contact_time"])
    return result


def format_duration(value: Any, language: str) -> str:
    text = clean_text(value)
    if not text:
        return ui(language, "not_specified")
    match = re.fullmatch(r"P(?:(\d+)Y)?(?:(\d+)M)?(?:(\d+)D)?", text)
    if not match:
        return text
    years, months, days = (int(part or 0) for part in match.groups())
    parts: list[str] = []
    if years:
        label = ui(language, "year_singular") if years == 1 else ui(language, "years")
        parts.append(f"{years} {label}")
    if months:
        label = ui(language, "month") if months == 1 else ui(language, "months")
        parts.append(f"{months} {label}")
    if days:
        day_label = "Tag" if base_language(language) == "de" and days == 1 else (
            "Tage" if base_language(language) == "de" else ("day" if days == 1 else "days")
        )
        parts.append(f"{days} {day_label}")
    return ", ".join(parts) or text


def language_list(values: Iterable[Any], language: str) -> str:
    lang = base_language(language)
    labels = []
    for item in values or []:
        code = str(item).lower()
        labels.append(LANGUAGE_LABELS.get(code, {}).get(lang, str(item)))
    return ", ".join(dict.fromkeys(labels)) or ui(language, "not_specified")


def delivery_list(values: Iterable[Any], language: str) -> str:
    lang = base_language(language)
    labels = []
    for item in values or []:
        key = str(item).lower()
        labels.append(DELIVERY_LABELS.get(key, {}).get(lang, str(item).replace("_", " ").title()))
    return ", ".join(dict.fromkeys(labels)) or ui(language, "not_specified")


def human_level(value: Any, language: str) -> str:
    text = clean_text(value)
    if not text:
        return ui(language, "not_specified")
    mapping_en = {"bachelor": "Bachelor (EQF 6)", "master": "Master (EQF 7)", "doctoral": "Doctoral (EQF 8)"}
    mapping_de = {"bachelor": "Bachelor (EQR 6)", "master": "Master (EQR 7)", "doctoral": "Promotion (EQR 8)"}
    mapping = mapping_de if base_language(language) == "de" else mapping_en
    return mapping.get(text.lower(), text.replace("_", " ").title())


def complexity_label(value: Any, language: str) -> str:
    key = clean_text(value).lower()
    return COMPLEXITY_LABELS.get(key, {}).get(base_language(language), key.replace("_", " ").title())


def safe_bookmark_name(code: str) -> str:
    name = re.sub(r"[^A-Za-z0-9_]", "_", code)
    if not name or name[0].isdigit():
        name = "M_" + name
    return name[:38]


@dataclass
class CodebookIndex:
    domains: dict[str, dict[str, Any]]
    outcomes_by_id: dict[str, dict[str, Any]]
    outcomes_by_code: dict[str, dict[str, Any]]

    @classmethod
    def empty(cls) -> "CodebookIndex":
        return cls({}, {}, {})

    @classmethod
    def from_json(cls, data: Mapping[str, Any]) -> "CodebookIndex":
        domains: dict[str, dict[str, Any]] = {}
        by_id: dict[str, dict[str, Any]] = {}
        by_code: dict[str, dict[str, Any]] = {}
        for entry in data.get("learningOutcomes", []) or []:
            if not isinstance(entry, Mapping):
                continue
            entry_dict = dict(entry)
            entry_id = clean_text(entry.get("learningOutcomeId"))
            primary = entry.get("primaryCode") or {}
            code = clean_text(primary.get("code")) if isinstance(primary, Mapping) else ""
            code_type = clean_text(primary.get("codeType")) if isinstance(primary, Mapping) else ""
            if entry_id:
                by_id[entry_id] = entry_dict
            if code:
                by_code[code] = entry_dict
            if code_type == "x-dhbw-domain" and code:
                domains[code] = entry_dict
        return cls(domains, by_id, by_code)

    def enrich_outcome(self, outcome: Mapping[str, Any]) -> dict[str, Any]:
        """Keep module-specific wording; fill only fields that are absent."""
        result = dict(outcome)
        primary = outcome.get("primaryCode") or {}
        code = clean_text(primary.get("code")) if isinstance(primary, Mapping) else ""
        outcome_id = clean_text(outcome.get("learningOutcomeId"))
        source = self.outcomes_by_id.get(outcome_id) or self.outcomes_by_code.get(code) or {}
        for field in ("name", "description", "complexityLevel", "primaryCode"):
            if not result.get(field) and source.get(field):
                result[field] = source[field]
        return result


@dataclass
class ModuleRecord:
    path: Path
    data: dict[str, Any]
    code: str
    title: str
    stream: str
    academic_year: int | None
    ects: float | None


def module_record(path: Path, data: dict[str, Any], language: str) -> ModuleRecord:
    code = clean_text(data.get("code")) or path.stem
    title = localized_value(data.get("name"), language) or code
    ext = data.get("ext") or {}
    stream = clean_text(ext.get("x-dhbw-stream")) if isinstance(ext, Mapping) else ""
    if not stream and isinstance(ext, Mapping):
        new_study = ext.get("newStudy") or {}
        if isinstance(new_study, Mapping):
            stream = clean_text(new_study.get("stream"))
    year_value = None
    if isinstance(ext, Mapping):
        year_value = ext.get("x-dhbw-academic-year")
        if year_value is None and isinstance(ext.get("newStudy"), Mapping):
            year_value = ext["newStudy"].get("year")
    try:
        academic_year = int(year_value) if year_value not in (None, "") else None
    except (TypeError, ValueError):
        academic_year = None
    loads = study_load_map(data)
    return ModuleRecord(path, data, code, title, stream, academic_year, loads.get("ects"))


# ---------------------------------------------------------------------------
# Word/OOXML helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_width(table, widths) -> None:
    """Apply exact column widths and prevent Word from expanding the table.

    ``python-docx`` column widths alone are sometimes treated as suggestions by
    Word.  Setting fixed layout plus the width on every cell keeps wide index
    tables inside the printable page area.
    """
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = width
        for cell in table.columns[col_idx].cells:
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")


def remove_paragraph_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, *, bold: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    r_pr.append(color)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_borders(table, color: str = "D7E0E2", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_label_value(cell, label: str, value: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(label.upper())
    set_run_font(label_run, size=7.5, bold=True, color=MID_GREY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(0)
    value_run = p2.add_run(value or "—")
    set_run_font(value_run, size=9.5, bold=True, color=DARK_GREY)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)


def add_heading_bar(document: Document, text: str, level: int = 2):
    p = document.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def set_document_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------
def configure_document(document: Document, title: str, institution: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 30, DARK_TEAL),
        ("Subtitle", 13, MID_GREY),
        ("Heading 1", 20, DARK_TEAL),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 11, DARK_TEAL),
    ):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
        if style_name == "Title":
            style_p_pr = style._element.get_or_add_pPr()
            style_border = style_p_pr.find(qn("w:pBdr"))
            if style_border is not None:
                style_p_pr.remove(style_border)

    # List style font consistency
    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10)

    props = document.core_properties
    props.title = title
    props.subject = "Human-readable module handbook assembled from OOAPI-style JSON"
    props.author = institution
    props.keywords = "module handbook, curriculum, learning outcomes, OOAPI"

    # Header and footer
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{institution}  ·  {title}")
    set_run_font(run, size=8, bold=True, color=TEAL)

    footer = section.footer
    fp = footer.paragraphs[0]
    left_run = fp.add_run(institution)
    set_run_font(left_run, size=7.5, color=MID_GREY)
    fp.add_run("    ")
    add_page_number(fp)

    set_document_update_fields(document)


def add_cover(
    document: Document,
    *,
    title: str,
    programme: str,
    institution: str,
    assembled_date: str,
    records: Sequence[ModuleRecord],
    language: str,
) -> None:
    # Top accent line
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_cell_shading(table.cell(0, 0), TEAL)
    table.cell(0, 0).height = Cm(0.18)
    remove_table_borders(table)

    for _ in range(4):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(institution)
    set_run_font(r, size=12, bold=True, color=TEAL)

    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_borders(p)
    p.add_run(title)

    if programme:
        p = document.add_paragraph(style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(programme)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(ui(language, "source_note"))
    set_run_font(r, size=10.5, italic=True, color=MID_GREY)

    document.add_paragraph()

    loads = [record.ects for record in records if record.ects is not None]
    total_ects = sum(loads) if loads else None
    streams = {record.stream for record in records if record.stream}
    years = {record.academic_year for record in records if record.academic_year is not None}

    summary = document.add_table(rows=1, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.columns[0].width = Cm(3.6)
    summary.columns[1].width = Cm(3.6)
    summary.columns[2].width = Cm(3.6)
    summary.columns[3].width = Cm(3.6)
    labels_values = [
        (ui(language, "modules"), str(len(records))),
        (ui(language, "total_ects"), pretty_number(total_ects) if total_ects is not None else "—"),
        (ui(language, "streams"), str(len(streams))),
        (ui(language, "academic_years"), str(len(years))),
    ]
    for cell, (label, value) in zip(summary.rows[0].cells, labels_values):
        set_cell_shading(cell, PALE_TEAL)
        add_label_value(cell, label, value)
        cell.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(summary, color="C8D9DA", size="6")

    for _ in range(5):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{ui(language, 'assembled')}: {assembled_date}")
    set_run_font(r, size=9, color=MID_GREY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ui(language, "generated"))
    set_run_font(r, size=8.5, italic=True, color=MID_GREY)

    document.add_page_break()


def add_module_index(document: Document, records: Sequence[ModuleRecord], language: str) -> None:
    p = document.add_paragraph(style="Heading 1")
    p.add_run(ui(language, "module_index"))

    intro = document.add_paragraph()
    intro.add_run(
        "Click a module title to jump to its full description."
        if base_language(language) == "en"
        else "Klicken Sie auf einen Modultitel, um direkt zur vollständigen Beschreibung zu springen."
    )

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Keep the index comfortably inside the 17 cm printable width of an A4
    # page with 2 cm margins.  The previous 16.9 cm grid left no tolerance for
    # Word's cell padding and border calculations and could spill into the
    # page edge in some Windows/Word versions.
    widths = [Cm(2.0), Cm(6.5), Cm(4.0), Cm(1.5), Cm(1.5)]  # 15.5 cm total
    set_fixed_table_width(table, widths)
    headers = [ui(language, key) for key in ("code", "module", "stream", "year", "ects")]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TEAL)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=8.3, bold=True, color=WHITE)
        set_cell_margins(cell, top=80, start=75, bottom=80, end=75)
    set_repeat_table_header(table.rows[0])

    for idx, record in enumerate(records):
        cells = table.add_row().cells
        fill = WHITE if idx % 2 == 0 else PALE_GREY
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=80, start=75, bottom=80, end=75)
        for cell in cells:
            cell.text = ""

        p = cells[0].paragraphs[0]
        r = p.add_run(record.code)
        set_run_font(r, size=8.5, bold=True, color=DARK_TEAL)

        p = cells[1].paragraphs[0]
        add_internal_hyperlink(p, record.title, safe_bookmark_name(record.code), bold=True)

        values = [
            record.stream or "—",
            str(record.academic_year) if record.academic_year is not None else "—",
            pretty_number(record.ects) if record.ects is not None else "—",
        ]
        for cell, value in zip(cells[2:], values):
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=8.5, color=DARK_GREY)

    set_table_borders(table)
    document.add_page_break()


def add_module_title(document: Document, record: ModuleRecord, bookmark_id: int) -> None:
    p = document.add_paragraph(style="Heading 1")
    code_run = p.add_run(record.code)
    set_run_font(code_run, size=11, bold=True, color=ACCENT)
    sep = p.add_run("  ·  ")
    set_run_font(sep, size=11, color=MID_GREY)
    title_run = p.add_run(record.title)
    set_run_font(title_run, size=20, bold=True, color=DARK_TEAL)
    add_bookmark(p, safe_bookmark_name(record.code), bookmark_id)

    if record.stream:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(record.stream)
        set_run_font(r, size=9.5, italic=True, color=TEAL)


def at_a_glance_values(record: ModuleRecord, language: str) -> list[tuple[str, str]]:
    module = record.data
    ext = module.get("ext") or {}
    loads = study_load_map(module)
    eqf = clean_text(ext.get("x-dhbw-eqf-level")) if isinstance(ext, Mapping) else ""
    level = human_level(module.get("level"), language)
    if eqf and eqf not in level:
        level = f"{level} · EQF {eqf}" if base_language(language) == "en" else f"{level} · EQR {eqf}"

    assessment = localized_value(module.get("assessment"), language) or ui(language, "not_specified")
    values = [
        (ui(language, "ects"), pretty_number(loads.get("ects")) or "—"),
        (ui(language, "workload"), f"{pretty_number(loads.get('sbu'))} {ui(language, 'hours')}" if loads.get("sbu") is not None else "—"),
        (ui(language, "contact"), f"{pretty_number(loads.get('contact_time'))} {ui(language, 'hours')}" if loads.get("contact_time") is not None else "—"),
        (ui(language, "self_study"), f"{pretty_number(loads.get('self_study'))} {ui(language, 'hours')}" if loads.get("self_study") is not None else "—"),
        (ui(language, "year"), str(record.academic_year) if record.academic_year is not None else "—"),
        (ui(language, "level"), level),
        (ui(language, "duration"), format_duration(module.get("duration"), language)),
        (ui(language, "delivery"), delivery_list(module.get("modesOfDelivery", []), language)),
        (ui(language, "languages"), language_list(module.get("languages", []), language)),
        (ui(language, "assessment"), assessment),
    ]
    return values


def add_at_a_glance(document: Document, record: ModuleRecord, language: str) -> None:
    p = document.add_paragraph(style="Heading 2")
    p.add_run(ui(language, "at_a_glance"))

    values = at_a_glance_values(record, language)
    rows = (len(values) + 1) // 2
    table = document.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8.3)
    table.columns[1].width = Cm(8.3)
    for idx, (label, value) in enumerate(values):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        set_cell_shading(cell, PALE_TEAL if row % 2 == 0 else WHITE)
        add_label_value(cell, label, value)
    if len(values) % 2:
        set_cell_shading(table.cell(rows - 1, 1), PALE_TEAL if (rows - 1) % 2 == 0 else WHITE)
    set_table_borders(table, color="C8D9DA", size="5")


def add_overview(document: Document, record: ModuleRecord, language: str) -> None:
    text = localized_value(record.data.get("description"), language)
    if not text:
        return
    p = document.add_paragraph(style="Heading 2")
    p.add_run(ui(language, "overview"))
    for block in text.split("\n"):
        if block.strip():
            document.add_paragraph(block.strip())


def domain_for_outcome(outcome: Mapping[str, Any]) -> str:
    primary = outcome.get("primaryCode") or {}
    code = clean_text(primary.get("code")) if isinstance(primary, Mapping) else ""
    if "-" in code:
        return code.split("-", 1)[0]
    return code or "OTHER"


def add_learning_outcomes(
    document: Document,
    record: ModuleRecord,
    codebook: CodebookIndex,
    language: str,
) -> None:
    raw_outcomes = record.data.get("learningOutcomes", []) or []
    outcomes = [codebook.enrich_outcome(item) for item in raw_outcomes if isinstance(item, Mapping)]
    if not outcomes:
        return

    p = document.add_paragraph(style="Heading 2")
    p.add_run(ui(language, "learning_outcomes"))

    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for outcome in outcomes:
        groups[domain_for_outcome(outcome)].append(outcome)

    def group_key(item: tuple[str, list[dict[str, Any]]]) -> tuple[str, str]:
        code, _ = item
        domain = codebook.domains.get(code, {})
        return (localized_value(domain.get("name"), language) or code, code)

    for domain_code, group in sorted(groups.items(), key=group_key):
        domain = codebook.domains.get(domain_code, {})
        domain_name = localized_value(domain.get("name"), language) or domain_code.replace("_", " ").title()
        domain_description = localized_value(domain.get("description"), language)

        p = document.add_paragraph(style="Heading 3")
        r = p.add_run(domain_name)
        set_run_font(r, size=11, bold=True, color=DARK_TEAL)
        if domain_description:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(domain_description)
            set_run_font(r, size=8.8, italic=True, color=MID_GREY)

        group.sort(
            key=lambda outcome: natural_sort_key(
                clean_text(((outcome.get("primaryCode") or {}).get("code")))
            )
        )
        for outcome in group:
            primary = outcome.get("primaryCode") or {}
            code = clean_text(primary.get("code")) if isinstance(primary, Mapping) else ""
            name = localized_value(outcome.get("name"), language) or code
            description = localized_value(outcome.get("description"), language)
            level = complexity_label(outcome.get("complexityLevel"), language)

            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(name)
            set_run_font(r, size=10, bold=True, color=DARK_GREY)
            if level:
                r = p.add_run(f"  ·  {level}")
                set_run_font(r, size=8.3, italic=True, color=TEAL)
            if description:
                p2 = document.add_paragraph()
                p2.paragraph_format.left_indent = Cm(0.85)
                p2.paragraph_format.space_after = Pt(5)
                r = p2.add_run(description)
                set_run_font(r, size=9.5, color=DARK_GREY)


def add_resources(document: Document, record: ModuleRecord, language: str) -> None:
    resources = [clean_text(item) for item in (record.data.get("resources", []) or []) if clean_text(item)]
    if not resources:
        return
    p = document.add_paragraph(style="Heading 2")
    p.add_run(ui(language, "resources"))
    for item in resources:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        r = p.add_run(item)
        set_run_font(r, size=9.4, color=DARK_GREY)


def add_source_note(document: Document, record: ModuleRecord, language: str, include_source: bool) -> None:
    if not include_source:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f"{ui(language, 'source')}: {record.path.name}")
    set_run_font(r, size=7.5, italic=True, color=MID_GREY)


def build_handbook(
    records: Sequence[ModuleRecord],
    codebook: CodebookIndex,
    output_path: Path,
    *,
    title: str,
    programme: str,
    institution: str,
    language: str,
    assembled_date: str,
    include_source_filenames: bool,
) -> None:
    document = Document()
    configure_document(document, title, institution)
    add_cover(
        document,
        title=title,
        programme=programme,
        institution=institution,
        assembled_date=assembled_date,
        records=records,
        language=language,
    )
    add_module_index(document, records, language)

    bookmark_id = 100
    for idx, record in enumerate(records):
        if idx:
            document.add_page_break()
        add_module_title(document, record, bookmark_id)
        bookmark_id += 1
        add_at_a_glance(document, record, language)
        add_overview(document, record, language)
        add_learning_outcomes(document, record, codebook, language)
        add_resources(document, record, language)
        add_source_note(document, record, language, include_source_filenames)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


# ---------------------------------------------------------------------------
# Markdown building
# ---------------------------------------------------------------------------
def markdown_anchor(code: str) -> str:
    """Return a stable explicit HTML anchor for a module code."""
    anchor = re.sub(r"[^A-Za-z0-9_-]+", "-", clean_text(code)).strip("-").lower()
    return f"module-{anchor or 'unknown'}"


def markdown_table_text(value: Any) -> str:
    """Escape content used inside a Markdown table cell."""
    text = clean_text(value).replace("\\", "\\\\").replace("|", "\\|")
    return text.replace("\n", "<br>") or "—"


def markdown_inline_text(value: Any) -> str:
    """Apply minimal escaping for Markdown link labels and bold text."""
    return clean_text(value).replace("\\", "\\\\").replace("[", "\\[").replace("]", "\\]")


def markdown_paragraphs(text: str) -> list[str]:
    return [block.strip() for block in clean_text(text).split("\n") if block.strip()]


def build_markdown_handbook(
    records: Sequence[ModuleRecord],
    codebook: CodebookIndex,
    output_path: Path,
    *,
    title: str,
    programme: str,
    institution: str,
    language: str,
    assembled_date: str,
    include_source_filenames: bool,
) -> None:
    """Write a clean, navigable Markdown edition of the module handbook."""
    lines: list[str] = []
    lines.append(f"# {markdown_inline_text(title)}")
    lines.append("")
    lines.append(f"**{markdown_inline_text(institution)}**")
    if programme:
        lines.append("")
        lines.append(f"*{markdown_inline_text(programme)}*")
    lines.append("")
    lines.append(f"> {ui(language, 'source_note')}")
    lines.append("")
    lines.append(f"**{ui(language, 'assembled')}:** {assembled_date}")
    lines.append("")

    loads = [record.ects for record in records if record.ects is not None]
    total_ects = sum(loads) if loads else None
    streams = {record.stream for record in records if record.stream}
    years = {record.academic_year for record in records if record.academic_year is not None}

    lines.append(f"## {ui(language, 'collection_overview')}")
    lines.append("")
    lines.append(
        f"| {ui(language, 'modules')} | {ui(language, 'total_ects')} | "
        f"{ui(language, 'streams')} | {ui(language, 'academic_years')} |"
    )
    lines.append("|---:|---:|---:|---:|")
    lines.append(
        f"| {len(records)} | {pretty_number(total_ects) if total_ects is not None else '—'} | "
        f"{len(streams)} | {len(years)} |"
    )
    lines.append("")

    lines.append(f"## {ui(language, 'module_index')}")
    lines.append("")
    lines.append(
        "Click a module title to jump to its full description."
        if base_language(language) == "en"
        else "Klicken Sie auf einen Modultitel, um direkt zur vollständigen Beschreibung zu springen."
    )
    lines.append("")
    lines.append(
        f"| {ui(language, 'code')} | {ui(language, 'module')} | {ui(language, 'stream')} | "
        f"{ui(language, 'year')} | {ui(language, 'ects')} |"
    )
    lines.append("|---|---|---|---:|---:|")
    for record in records:
        title_link = (
            f"[{markdown_inline_text(record.title)}](#{markdown_anchor(record.code)})"
        )
        lines.append(
            f"| {markdown_table_text(record.code)} | {title_link} | "
            f"{markdown_table_text(record.stream or '—')} | "
            f"{record.academic_year if record.academic_year is not None else '—'} | "
            f"{pretty_number(record.ects) if record.ects is not None else '—'} |"
        )
    lines.append("")

    for record in records:
        lines.append("---")
        lines.append("")
        lines.append(f'<a id="{markdown_anchor(record.code)}"></a>')
        lines.append("")
        lines.append(
            f"## {markdown_inline_text(record.code)} · {markdown_inline_text(record.title)}"
        )
        lines.append("")
        if record.stream:
            lines.append(f"*{markdown_inline_text(record.stream)}*")
            lines.append("")

        lines.append(f"### {ui(language, 'at_a_glance')}")
        lines.append("")
        lines.append("| | |")
        lines.append("|---|---|")
        for label, value in at_a_glance_values(record, language):
            lines.append(
                f"| **{markdown_table_text(label)}** | {markdown_table_text(value)} |"
            )
        lines.append("")

        overview = localized_value(record.data.get("description"), language)
        if overview:
            lines.append(f"### {ui(language, 'overview')}")
            lines.append("")
            for paragraph in markdown_paragraphs(overview):
                lines.append(paragraph)
                lines.append("")

        raw_outcomes = record.data.get("learningOutcomes", []) or []
        outcomes = [
            codebook.enrich_outcome(item)
            for item in raw_outcomes
            if isinstance(item, Mapping)
        ]
        if outcomes:
            lines.append(f"### {ui(language, 'learning_outcomes')}")
            lines.append("")
            groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
            for outcome in outcomes:
                groups[domain_for_outcome(outcome)].append(outcome)

            def group_key(item: tuple[str, list[dict[str, Any]]]) -> tuple[str, str]:
                code, _ = item
                domain = codebook.domains.get(code, {})
                return (localized_value(domain.get("name"), language) or code, code)

            for domain_code, group in sorted(groups.items(), key=group_key):
                domain = codebook.domains.get(domain_code, {})
                domain_name = (
                    localized_value(domain.get("name"), language)
                    or domain_code.replace("_", " ").title()
                )
                domain_description = localized_value(domain.get("description"), language)
                lines.append(f"#### {markdown_inline_text(domain_name)}")
                lines.append("")
                if domain_description:
                    lines.append(f"*{domain_description}*")
                    lines.append("")

                group.sort(
                    key=lambda outcome: natural_sort_key(
                        clean_text(((outcome.get("primaryCode") or {}).get("code")))
                    )
                )
                for outcome in group:
                    primary = outcome.get("primaryCode") or {}
                    code = clean_text(primary.get("code")) if isinstance(primary, Mapping) else ""
                    name = localized_value(outcome.get("name"), language) or code
                    description = localized_value(outcome.get("description"), language)
                    level = complexity_label(outcome.get("complexityLevel"), language)
                    label = f"**{markdown_inline_text(name)}**"
                    if level:
                        label += f" — *{markdown_inline_text(level)}*"
                    lines.append(f"- {label}")
                    if description:
                        lines.append(f"  {description}")
                lines.append("")

        resources = [
            clean_text(item)
            for item in (record.data.get("resources", []) or [])
            if clean_text(item)
        ]
        if resources:
            lines.append(f"### {ui(language, 'resources')}")
            lines.append("")
            for item in resources:
                lines.append(f"- {item}")
            lines.append("")

        if include_source_filenames:
            lines.append(f"*{ui(language, 'source')}: `{record.path.name}`*")
            lines.append("")

        lines.append(f"[↑ {ui(language, 'module_index')}](#module-index)")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append(f"*{ui(language, 'generated')}*")
    lines.append("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert OOAPI-style module JSON files into readable Word and Markdown handbooks."
    )
    parser.add_argument("--dir", required=True, type=Path, help="Directory containing module JSON files")
    parser.add_argument(
        "--pattern",
        default="T*INF*v6.json",
        help="Glob used to select module files (default: T*INF*v6.json)",
    )
    parser.add_argument("--codebook", type=Path, help="Optional competence codebook JSON")
    parser.add_argument("--output", required=True, type=Path, help="Output .docx path")
    parser.add_argument(
        "--markdown-output",
        type=Path,
        help="Optional output .md path (default: same name as --output with .md suffix)",
    )
    parser.add_argument("--language", default="en-GB", choices=("en-GB", "de-DE"))
    parser.add_argument("--title", default="Module Handbook")
    parser.add_argument("--programme", default="Bachelor of Computer Science (B.Sc.)")
    parser.add_argument("--institution", default="Open University of Germany (OUG)")
    parser.add_argument(
        "--date",
        dest="assembled_date",
        default=date.today().isoformat(),
        help="Assembly date shown on the cover (default: today, YYYY-MM-DD)",
    )
    parser.add_argument(
        "--include-source-filenames",
        action="store_true",
        help="Show each JSON filename in small text at the end of its module",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    input_dir: Path = args.dir
    if not input_dir.is_dir():
        print(f"ERROR: directory does not exist: {input_dir}", file=sys.stderr)
        return 2
    if args.output.suffix.lower() != ".docx":
        print("ERROR: --output must end in .docx", file=sys.stderr)
        return 2

    markdown_output: Path = args.markdown_output or args.output.with_suffix(".md")
    if markdown_output.suffix.lower() not in {".md", ".markdown"}:
        print("ERROR: --markdown-output must end in .md or .markdown", file=sys.stderr)
        return 2

    try:
        datetime.fromisoformat(args.assembled_date)
    except ValueError:
        print("ERROR: --date must be an ISO date such as 2026-07-30", file=sys.stderr)
        return 2

    paths = sorted(input_dir.glob(args.pattern), key=lambda path: natural_sort_key(path.name))
    if not paths:
        print(
            f"ERROR: no files matched {args.pattern!r} in {input_dir.resolve()}",
            file=sys.stderr,
        )
        return 2

    records: list[ModuleRecord] = []
    seen_codes: dict[str, Path] = {}
    errors: list[str] = []
    for path in paths:
        try:
            data = load_json(path)
            record = module_record(path, data, args.language)
            if record.code in seen_codes:
                raise ValueError(
                    f"duplicate module code {record.code!r}; already found in {seen_codes[record.code].name}"
                )
            seen_codes[record.code] = path
            records.append(record)
        except Exception as exc:  # report all input errors together
            errors.append(f"{path.name}: {exc}")

    if errors:
        print("ERROR: one or more module files could not be processed:", file=sys.stderr)
        for error in errors:
            print(f"  - {error}", file=sys.stderr)
        return 2

    records.sort(
        key=lambda record: (
            record.academic_year if record.academic_year is not None else 999,
            record.stream.lower() if record.stream else "zzzz",
            natural_sort_key(record.code),
        )
    )

    codebook = CodebookIndex.empty()
    if args.codebook:
        if not args.codebook.is_file():
            print(f"ERROR: codebook does not exist: {args.codebook}", file=sys.stderr)
            return 2
        try:
            codebook = CodebookIndex.from_json(load_json(args.codebook))
        except Exception as exc:
            print(f"ERROR: codebook could not be read: {exc}", file=sys.stderr)
            return 2

    build_handbook(
        records,
        codebook,
        args.output,
        title=args.title,
        programme=args.programme,
        institution=args.institution,
        language=args.language,
        assembled_date=args.assembled_date,
        include_source_filenames=args.include_source_filenames,
    )
    build_markdown_handbook(
        records,
        codebook,
        markdown_output,
        title=args.title,
        programme=args.programme,
        institution=args.institution,
        language=args.language,
        assembled_date=args.assembled_date,
        include_source_filenames=args.include_source_filenames,
    )

    print(f"Created {args.output.resolve()}")
    print(f"Created {markdown_output.resolve()}")
    print(f"Modules: {len(records)}")
    print(f"Pattern: {args.pattern}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
