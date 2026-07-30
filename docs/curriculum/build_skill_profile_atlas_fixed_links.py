#!/usr/bin/env python3
r"""Build a visual skill-profile atlas from OOAPI-style module JSON files.

The script scans a directory for module files (default: ``T*INF*v6.json``),
enriches each module learning outcome from the DHBW/OUG competence codebook,
and creates two human-readable outputs from the same analysis:

* a designed Word report (DOCX), and
* a linked Markdown report with the same charts.

The report includes:

* portfolio summary and module index;
* competence-domain inventory;
* Bloom taxonomy distribution and progression by year;
* AI-skill inventory and AI-domain coverage;
* human-skill tags, meta-skills, and capability clusters;
* comparisons by year, semester (where available), and stream;
* a domain-by-stream coverage matrix;
* compact skill profiles for every selected module.

Machine-oriented material such as UUIDs, raw extension tags, legacy IDs,
processing notes, and source mappings is used only for analysis and is not
printed as technical clutter.

Example (PowerShell)
--------------------
python .\build_skill_profile_atlas.py `
  --dir ".\dir" `
  --codebook ".\competence_codebook_v11_ooapi.json" `
  --output ".\OUG_Skill_Profile_Atlas.docx" `
  --language "en-GB"

Dependencies
------------
python -m pip install python-docx matplotlib
"""

from __future__ import annotations

import argparse
import json
import math
import re
import sys
import textwrap
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "matplotlib is required. Install it with: python -m pip install matplotlib"
    ) from exc

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install it with: python -m pip install python-docx"
    ) from exc


# ---------------------------------------------------------------------------
# Visual system
# ---------------------------------------------------------------------------
TEAL = "2F5D62"
DARK_TEAL = "173F43"
PALE_TEAL = "EAF2F2"
PALE_GREY = "F4F6F7"
MID_GREY = "667276"
DARK_GREY = "263238"
WHITE = "FFFFFF"
ACCENT = "D8A43B"
FONT = "Arial"

BLOOM_ORDER = ["remember", "understand", "apply", "analyze", "evaluate", "create"]
BLOOM_LEVEL = {
    "remember": 1,
    "understand": 2,
    "apply": 2,
    "analyze": 3,
    "analyse": 3,
    "evaluate": 4,
    "create": 4,
}

UI = {
    "en": {
        "title": "Skill Profile Atlas",
        "subtitle": "Competence coverage, progression, AI readiness and human capabilities",
        "assembled": "Assembled",
        "source_note": "Generated from machine-readable module and competence data.",
        "portfolio_snapshot": "Portfolio snapshot",
        "modules": "Modules",
        "total_ects": "Total ECTS",
        "unique_skills": "Unique skills",
        "domains": "Domains touched",
        "ai_skills": "AI-related skills",
        "human_skills": "Human-tagged skills",
        "contents": "Report contents",
        "module_index": "Module index",
        "executive_profile": "Executive skill profile",
        "domain_inventory": "Competence-domain inventory",
        "bloom_profile": "Bloom taxonomy profile",
        "ai_profile": "AI skill profile",
        "human_profile": "Human and meta-skill profile",
        "by_year": "Progression by academic year",
        "by_semester": "Progression by semester",
        "by_stream": "Profiles by stream",
        "coverage_matrix": "Domain coverage across streams",
        "module_profiles": "Module skill profiles",
        "code": "Code",
        "module": "Module",
        "stream": "Stream",
        "year": "Year",
        "semester": "Semester",
        "semester_short": "Sem.",
        "ects": "ECTS",
        "skills": "Skills",
        "ai": "AI",
        "human": "Human",
        "domain": "Domain",
        "domain_name": "Domain name",
        "modules_touched": "Modules",
        "occurrences": "LO occurrences",
        "years": "Years",
        "streams": "Streams",
        "bloom": "Bloom level",
        "skill": "Skill",
        "ai_domains": "AI domains",
        "human_tag": "Human tag",
        "meta_category": "Meta-skill category",
        "clusters": "Capability clusters",
        "coverage": "Coverage",
        "not_available": "Not available",
        "no_data": "No matching data was found.",
        "interpretation": "How to read this profile",
        "interpretation_text": (
            "Counts show how often a competence appears across the selected modules. "
            "Unique-skill counts remove repetition, while learning-outcome occurrences "
            "show curricular emphasis. Bloom levels describe the cognitive demand encoded "
            "in the learning outcomes; they are not grades or performance results."
        ),
        "ai_interpretation": (
            "AI-related skills are identified from the codebook's AI flag, AI-domain labels, "
            "or AI-upskilling meta-skill marker. The report distinguishes the breadth of AI "
            "domains from repeated use of the same competence across modules."
        ),
        "human_interpretation": (
            "Human capability is represented through explicit human-skill tags, meta-skill "
            "categories and broader capability clusters. Empty fields are not interpreted as "
            "absence of human learning; they indicate that no explicit tag was present."
        ),
        "source": "Source file",
        "back": "Back to module index",
    },
    "de": {
        "title": "Kompetenzprofil-Atlas",
        "subtitle": "Kompetenzabdeckung, Progression, KI-Befähigung und menschliche Fähigkeiten",
        "assembled": "Zusammengestellt am",
        "source_note": "Erzeugt aus maschinenlesbaren Modul- und Kompetenzdaten.",
        "portfolio_snapshot": "Portfolio auf einen Blick",
        "modules": "Module",
        "total_ects": "ECTS gesamt",
        "unique_skills": "Einzigartige Kompetenzen",
        "domains": "Berührte Domänen",
        "ai_skills": "KI-bezogene Kompetenzen",
        "human_skills": "Menschlich markierte Kompetenzen",
        "contents": "Inhalt des Berichts",
        "module_index": "Modulübersicht",
        "executive_profile": "Gesamtprofil der Kompetenzen",
        "domain_inventory": "Inventar der Kompetenzdomänen",
        "bloom_profile": "Bloom-Taxonomie-Profil",
        "ai_profile": "KI-Kompetenzprofil",
        "human_profile": "Menschliche und Meta-Kompetenzen",
        "by_year": "Progression nach Studienjahr",
        "by_semester": "Progression nach Semester",
        "by_stream": "Profile nach Studienströmen",
        "coverage_matrix": "Domänenabdeckung nach Studienströmen",
        "module_profiles": "Kompetenzprofile der Module",
        "code": "Code",
        "module": "Modul",
        "stream": "Studienstrom",
        "year": "Jahr",
        "semester": "Semester",
        "semester_short": "Sem.",
        "ects": "ECTS",
        "skills": "Kompetenzen",
        "ai": "KI",
        "human": "Menschlich",
        "domain": "Domäne",
        "domain_name": "Domänenname",
        "modules_touched": "Module",
        "occurrences": "LO-Vorkommen",
        "years": "Jahre",
        "streams": "Studienströme",
        "bloom": "Bloom-Stufe",
        "skill": "Kompetenz",
        "ai_domains": "KI-Domänen",
        "human_tag": "Human-Tag",
        "meta_category": "Meta-Kompetenzkategorie",
        "clusters": "Fähigkeitscluster",
        "coverage": "Abdeckung",
        "not_available": "Nicht verfügbar",
        "no_data": "Keine passenden Daten gefunden.",
        "interpretation": "So ist das Profil zu lesen",
        "interpretation_text": (
            "Die Häufigkeiten zeigen, wie oft eine Kompetenz in den ausgewählten Modulen "
            "vorkommt. Einzigartige Kompetenzen werden nur einmal gezählt; LO-Vorkommen "
            "zeigen dagegen die curriculare Gewichtung. Bloom-Stufen beschreiben die in den "
            "Lernergebnissen angelegte kognitive Anforderung, nicht Noten oder Leistung."
        ),
        "ai_interpretation": (
            "KI-Kompetenzen werden über das KI-Kennzeichen, KI-Domänen oder die Markierung "
            "AI-Upskilling im Codebook identifiziert. Der Bericht unterscheidet die Breite "
            "der KI-Domänen von wiederholten Kompetenzvorkommen in mehreren Modulen."
        ),
        "human_interpretation": (
            "Menschliche Fähigkeiten werden durch explizite Human-Tags, Meta-Kompetenzkategorien "
            "und breitere Fähigkeitscluster sichtbar. Leere Felder bedeuten nicht, dass kein "
            "menschliches Lernen stattfindet, sondern nur, dass kein explizites Tag vorlag."
        ),
        "source": "Quelldatei",
        "back": "Zurück zur Modulübersicht",
    },
}

BLOOM_LABELS = {
    "en": {
        "remember": "Remember / foundation",
        "understand": "Understand / apply",
        "apply": "Apply",
        "analyze": "Analyse",
        "evaluate": "Evaluate / create",
        "create": "Create",
        "unknown": "Unspecified",
    },
    "de": {
        "remember": "Erinnern / Grundlage",
        "understand": "Verstehen / Anwenden",
        "apply": "Anwenden",
        "analyze": "Analysieren",
        "evaluate": "Bewerten / Gestalten",
        "create": "Gestalten",
        "unknown": "Nicht angegeben",
    },
}


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------
def base_language(language: str) -> str:
    return "de" if language.lower().startswith("de") else "en"


def tr(language: str, key: str) -> str:
    return UI[base_language(language)][key]


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ").replace("\u200b", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def natural_sort_key(text: str) -> list[Any]:
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", text)]


def load_json(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in {path}: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError(f"Expected a JSON object in {path}")
    return data


def localized_value(value: Any, language: str) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return clean_text(value)
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, Mapping):
        return clean_text(value.get("value", ""))
    if not isinstance(value, Sequence):
        return clean_text(value)

    candidates = [item for item in value if isinstance(item, Mapping)]
    if not candidates:
        return ""
    requested = language.lower()
    requested_base = requested.split("-")[0]
    for item in candidates:
        if str(item.get("language", "")).lower() == requested:
            return clean_text(item.get("value", ""))
    for item in candidates:
        if str(item.get("language", "")).lower().split("-")[0] == requested_base:
            return clean_text(item.get("value", ""))
    for preferred in ("en-GB", "de-DE", "en", "de"):
        for item in candidates:
            lang = str(item.get("language", "")).lower()
            if lang == preferred.lower() or lang.split("-")[0] == preferred.lower().split("-")[0]:
                return clean_text(item.get("value", ""))
    return clean_text(candidates[0].get("value", ""))


def as_number(value: Any) -> float | None:
    try:
        if value is None or value == "":
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def pretty_number(value: float | int | None) -> str:
    if value is None:
        return ""
    value = float(value)
    return str(int(value)) if value.is_integer() else f"{value:.1f}"


def unique_sorted(values: Iterable[str]) -> list[str]:
    return sorted({clean_text(v) for v in values if clean_text(v)}, key=natural_sort_key)


def first_present(*values: Any) -> Any:
    for value in values:
        if value is not None and value != "":
            return value
    return None


def get_nested(mapping: Mapping[str, Any], *path: str) -> Any:
    current: Any = mapping
    for key in path:
        if not isinstance(current, Mapping):
            return None
        current = current.get(key)
    return current


def study_load(module: Mapping[str, Any], unit: str) -> float | None:
    for item in module.get("studyLoad", []) or []:
        if isinstance(item, Mapping) and str(item.get("studyLoadUnit", "")).lower() == unit.lower():
            return as_number(item.get("value"))
    return None


def normalise_bloom(value: Any) -> str:
    raw = clean_text(value).lower()
    aliases = {
        "analyse": "analyze",
        "analysis": "analyze",
        "remember/understand": "remember",
        "understanding": "understand",
        "evaluation": "evaluate",
    }
    raw = aliases.get(raw, raw)
    if raw in BLOOM_ORDER:
        return raw
    # Accept labels such as "Evaluate/Create (Mastery)" or "L3".
    if "create" in raw:
        return "create"
    if "evaluat" in raw or "bewert" in raw:
        return "evaluate"
    if "analy" in raw:
        return "analyze"
    if "apply" in raw or "anwend" in raw:
        return "apply"
    if "understand" in raw or "versteh" in raw:
        return "understand"
    if "remember" in raw or "foundation" in raw or "grundlage" in raw:
        return "remember"
    if raw in {"l1", "1"}:
        return "remember"
    if raw in {"l2", "2"}:
        return "understand"
    if raw in {"l3", "3"}:
        return "analyze"
    if raw in {"l4", "4"}:
        return "evaluate"
    return "unknown"


def bloom_label(value: str, language: str) -> str:
    return BLOOM_LABELS[base_language(language)].get(value, BLOOM_LABELS[base_language(language)]["unknown"])


def module_semester(module: Mapping[str, Any]) -> str:
    ext = module.get("ext") or {}
    new_study = ext.get("newStudy") or {}
    candidates = (
        module.get("semester"),
        module.get("term"),
        ext.get("x-dhbw-semester"),
        ext.get("x-dhbw-recommended-semester"),
        ext.get("x-dhbw-term"),
        new_study.get("semester"),
    )
    value = first_present(*candidates)
    if isinstance(value, Sequence) and not isinstance(value, str):
        return ", ".join(clean_text(v) for v in value if clean_text(v))
    return clean_text(value)


def module_year(module: Mapping[str, Any]) -> str:
    ext = module.get("ext") or {}
    new_study = ext.get("newStudy") or {}
    value = first_present(
        ext.get("x-dhbw-academic-year"),
        new_study.get("year"),
        module.get("academicYear"),
        module.get("year"),
    )
    return clean_text(value)


def module_stream(module: Mapping[str, Any]) -> str:
    ext = module.get("ext") or {}
    new_study = ext.get("newStudy") or {}
    return clean_text(
        first_present(ext.get("x-dhbw-stream"), new_study.get("stream"), module.get("stream"))
    )


def markdown_escape(text: Any) -> str:
    return clean_text(text).replace("|", "\\|").replace("\n", "<br>")


def markdown_anchor(code: str) -> str:
    """Return a stable explicit HTML anchor based only on the module code.

    GitHub/Jekyll heading slugs vary when headings contain punctuation, Unicode,
    or repeated separators.  Using an explicit anchor keeps links stable even
    when a module title changes.
    """
    anchor = re.sub(r"[^A-Za-z0-9_-]+", "-", clean_text(code)).strip("-").lower()
    return f"module-{anchor or 'unknown'}"


# ---------------------------------------------------------------------------
# Data model and codebook enrichment
# ---------------------------------------------------------------------------
@dataclass
class DomainInfo:
    code: str
    name: str
    description: str
    learning_outcome_id: str = ""


@dataclass
class CompetenceInfo:
    code: str
    name: str
    description: str
    bloom: str
    domain_code: str
    ai_related: bool
    ai_domains: tuple[str, ...]
    human_tag: str
    meta_skill: bool
    meta_category: str
    clusters: tuple[str, ...]
    meta_skills: tuple[str, ...]
    learning_outcome_id: str = ""


@dataclass
class SkillOccurrence:
    module_code: str
    module_title: str
    module_stream: str
    module_year: str
    module_semester: str
    module_ects: float
    skill_code: str
    skill_name: str
    skill_description: str
    domain_code: str
    domain_name: str
    bloom: str
    ai_related: bool
    ai_domains: tuple[str, ...]
    human_tag: str
    meta_skill: bool
    meta_category: str
    clusters: tuple[str, ...]
    meta_skills: tuple[str, ...]


@dataclass
class ModuleProfile:
    path: Path
    code: str
    title: str
    description: str
    stream: str
    year: str
    semester: str
    ects: float
    outcomes: list[SkillOccurrence] = field(default_factory=list)

    @property
    def unique_skill_codes(self) -> set[str]:
        return {item.skill_code for item in self.outcomes if item.skill_code}

    @property
    def domains(self) -> set[str]:
        return {item.domain_code for item in self.outcomes if item.domain_code}

    @property
    def ai_skills(self) -> set[str]:
        return {item.skill_code for item in self.outcomes if item.ai_related}

    @property
    def human_skills(self) -> set[str]:
        return {
            item.skill_code
            for item in self.outcomes
            if item.human_tag or item.meta_skill or item.meta_category
        }


class CodebookIndex:
    def __init__(self, data: Mapping[str, Any], language: str):
        self.language = language
        self.domains_by_code: dict[str, DomainInfo] = {}
        self.domains_by_id: dict[str, DomainInfo] = {}
        self.skills_by_code: dict[str, CompetenceInfo] = {}
        self.skills_by_id: dict[str, CompetenceInfo] = {}

        rows = data.get("learningOutcomes", []) or []
        for row in rows:
            if not isinstance(row, Mapping):
                continue
            primary = row.get("primaryCode") or {}
            code_type = clean_text(primary.get("codeType"))
            code = clean_text(primary.get("code"))
            if code_type == "x-dhbw-domain" and code:
                domain = DomainInfo(
                    code=code,
                    name=localized_value(row.get("name"), language) or code,
                    description=localized_value(row.get("description"), language),
                    learning_outcome_id=clean_text(row.get("learningOutcomeId")),
                )
                self.domains_by_code[code] = domain
                if domain.learning_outcome_id:
                    self.domains_by_id[domain.learning_outcome_id] = domain

        for row in rows:
            if not isinstance(row, Mapping):
                continue
            primary = row.get("primaryCode") or {}
            code_type = clean_text(primary.get("codeType"))
            code = clean_text(primary.get("code"))
            if code_type != "x-dhbw-competence-id" or not code:
                continue
            ext = row.get("ext") or {}
            parent_ids = row.get("parentIds") or []
            domain_code = ""
            for parent_id in parent_ids:
                domain = self.domains_by_id.get(clean_text(parent_id))
                if domain:
                    domain_code = domain.code
                    break
            if not domain_code:
                domain_code = code.split("-")[0] if "-" in code else code[:4]

            ai_domains = tuple(unique_sorted(ext.get("x-dhbw-ai_domain") or []))
            meta_skills = tuple(unique_sorted(ext.get("x-dhbw-meta-skills") or []))
            ai_related = bool(ext.get("x-dhbw-ai_related")) or bool(ai_domains) or any(
                item.lower() == "ai-upskilling" for item in meta_skills
            )
            competence = CompetenceInfo(
                code=code,
                name=localized_value(row.get("name"), language) or code,
                description=localized_value(row.get("description"), language),
                bloom=normalise_bloom(
                    first_present(row.get("complexityLevel"), ext.get("x-dhbw-bloom_category"))
                ),
                domain_code=domain_code,
                ai_related=ai_related,
                ai_domains=ai_domains,
                human_tag=clean_text(ext.get("x-dhbw-human_skill_tag")),
                meta_skill=bool(ext.get("x-dhbw-meta_skill")),
                meta_category=clean_text(ext.get("x-dhbw-meta_skill_category")),
                clusters=tuple(unique_sorted(ext.get("x-dhbw-clusters") or [])),
                meta_skills=meta_skills,
                learning_outcome_id=clean_text(row.get("learningOutcomeId")),
            )
            self.skills_by_code[code] = competence
            if competence.learning_outcome_id:
                self.skills_by_id[competence.learning_outcome_id] = competence

    def skill_for(self, outcome: Mapping[str, Any]) -> CompetenceInfo | None:
        primary = outcome.get("primaryCode") or {}
        code = clean_text(primary.get("code"))
        outcome_id = clean_text(outcome.get("learningOutcomeId"))
        return self.skills_by_code.get(code) or self.skills_by_id.get(outcome_id)

    def domain_for_code(self, code: str) -> DomainInfo:
        if code in self.domains_by_code:
            return self.domains_by_code[code]
        return DomainInfo(code=code, name=code or "Unclassified", description="")


def load_modules(
    directory: Path,
    pattern: str,
    codebook: CodebookIndex,
    language: str,
) -> list[ModuleProfile]:
    paths = sorted(directory.glob(pattern), key=lambda p: natural_sort_key(p.name))
    if not paths:
        raise FileNotFoundError(
            f"No module files matched {pattern!r} in {directory.resolve()}"
        )

    modules: list[ModuleProfile] = []
    errors: list[str] = []
    seen_codes: dict[str, Path] = {}

    for path in paths:
        try:
            data = load_json(path)
            code = clean_text(data.get("code")) or path.stem
            if code in seen_codes:
                raise ValueError(
                    f"Duplicate module code {code!r}: {seen_codes[code].name} and {path.name}"
                )
            seen_codes[code] = path
            ects = study_load(data, "ects") or 0.0
            profile = ModuleProfile(
                path=path,
                code=code,
                title=localized_value(data.get("name"), language) or code,
                description=localized_value(data.get("description"), language),
                stream=module_stream(data) or "Unassigned",
                year=module_year(data) or "Unspecified",
                semester=module_semester(data),
                ects=ects,
            )

            for outcome in data.get("learningOutcomes", []) or []:
                if not isinstance(outcome, Mapping):
                    continue
                primary = outcome.get("primaryCode") or {}
                skill_code = clean_text(primary.get("code"))
                codebook_skill = codebook.skill_for(outcome)
                if codebook_skill:
                    domain_code = codebook_skill.domain_code
                    domain = codebook.domain_for_code(domain_code)
                    skill_name = localized_value(outcome.get("name"), language) or codebook_skill.name
                    skill_description = (
                        localized_value(outcome.get("description"), language)
                        or codebook_skill.description
                    )
                    bloom = normalise_bloom(
                        first_present(outcome.get("complexityLevel"), codebook_skill.bloom)
                    )
                    occurrence = SkillOccurrence(
                        module_code=profile.code,
                        module_title=profile.title,
                        module_stream=profile.stream,
                        module_year=profile.year,
                        module_semester=profile.semester,
                        module_ects=profile.ects,
                        skill_code=skill_code or codebook_skill.code,
                        skill_name=skill_name or codebook_skill.name,
                        skill_description=skill_description,
                        domain_code=domain_code,
                        domain_name=domain.name,
                        bloom=bloom,
                        ai_related=codebook_skill.ai_related,
                        ai_domains=codebook_skill.ai_domains,
                        human_tag=codebook_skill.human_tag,
                        meta_skill=codebook_skill.meta_skill,
                        meta_category=codebook_skill.meta_category,
                        clusters=codebook_skill.clusters,
                        meta_skills=codebook_skill.meta_skills,
                    )
                else:
                    domain_code = skill_code.split("-")[0] if "-" in skill_code else "UNCL"
                    domain = codebook.domain_for_code(domain_code)
                    outcome_ext = outcome.get("ext") or {}
                    ai_domains = tuple(unique_sorted(outcome_ext.get("x-dhbw-ai_domain") or []))
                    meta_skills = tuple(unique_sorted(outcome_ext.get("x-dhbw-meta-skills") or []))
                    occurrence = SkillOccurrence(
                        module_code=profile.code,
                        module_title=profile.title,
                        module_stream=profile.stream,
                        module_year=profile.year,
                        module_semester=profile.semester,
                        module_ects=profile.ects,
                        skill_code=skill_code or clean_text(outcome.get("learningOutcomeId")),
                        skill_name=localized_value(outcome.get("name"), language) or skill_code,
                        skill_description=localized_value(outcome.get("description"), language),
                        domain_code=domain_code,
                        domain_name=domain.name,
                        bloom=normalise_bloom(outcome.get("complexityLevel")),
                        ai_related=bool(outcome_ext.get("x-dhbw-ai_related")) or bool(ai_domains),
                        ai_domains=ai_domains,
                        human_tag=clean_text(outcome_ext.get("x-dhbw-human_skill_tag")),
                        meta_skill=bool(outcome_ext.get("x-dhbw-meta_skill")),
                        meta_category=clean_text(outcome_ext.get("x-dhbw-meta_skill_category")),
                        clusters=tuple(unique_sorted(outcome_ext.get("x-dhbw-clusters") or [])),
                        meta_skills=meta_skills,
                    )
                profile.outcomes.append(occurrence)
            modules.append(profile)
        except Exception as exc:  # collect all input problems
            errors.append(f"{path.name}: {exc}")

    if errors:
        raise ValueError("Problems while reading module files:\n- " + "\n- ".join(errors))
    return modules


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------
@dataclass
class Analysis:
    modules: list[ModuleProfile]
    occurrences: list[SkillOccurrence]
    unique_skills: dict[str, SkillOccurrence]
    domain_occurrences: Counter[str]
    domain_unique_skills: dict[str, set[str]]
    domain_modules: dict[str, set[str]]
    domain_years: dict[str, set[str]]
    domain_streams: dict[str, set[str]]
    bloom_occurrences: Counter[str]
    ai_occurrences: list[SkillOccurrence]
    ai_unique_skills: dict[str, SkillOccurrence]
    ai_domain_counts: Counter[str]
    human_occurrences: list[SkillOccurrence]
    human_unique_skills: dict[str, SkillOccurrence]
    human_tag_counts: Counter[str]
    meta_category_counts: Counter[str]
    cluster_counts: Counter[str]
    year_groups: dict[str, list[ModuleProfile]]
    semester_groups: dict[str, list[ModuleProfile]]
    stream_groups: dict[str, list[ModuleProfile]]


def analyse(modules: list[ModuleProfile]) -> Analysis:
    occurrences = [outcome for module in modules for outcome in module.outcomes]
    unique_skills: dict[str, SkillOccurrence] = {}
    for item in occurrences:
        if item.skill_code and item.skill_code not in unique_skills:
            unique_skills[item.skill_code] = item

    domain_occurrences = Counter(item.domain_code for item in occurrences if item.domain_code)
    domain_unique_skills: dict[str, set[str]] = defaultdict(set)
    domain_modules: dict[str, set[str]] = defaultdict(set)
    domain_years: dict[str, set[str]] = defaultdict(set)
    domain_streams: dict[str, set[str]] = defaultdict(set)
    for item in occurrences:
        if item.domain_code:
            domain_unique_skills[item.domain_code].add(item.skill_code)
            domain_modules[item.domain_code].add(item.module_code)
            domain_years[item.domain_code].add(item.module_year)
            domain_streams[item.domain_code].add(item.module_stream)

    bloom_occurrences = Counter(item.bloom for item in occurrences)
    ai_occurrences = [item for item in occurrences if item.ai_related]
    ai_unique_skills = {item.skill_code: item for item in ai_occurrences if item.skill_code}
    ai_domain_counts = Counter(
        ai_domain
        for item in ai_occurrences
        for ai_domain in (item.ai_domains or ("AI-related (unspecified)",))
    )

    human_occurrences = [
        item
        for item in occurrences
        if item.human_tag or item.meta_skill or item.meta_category
    ]
    human_unique_skills = {item.skill_code: item for item in human_occurrences if item.skill_code}
    human_tag_counts = Counter(
        item.human_tag or "Human capability (unspecified)" for item in human_occurrences
    )
    meta_category_counts = Counter(
        item.meta_category for item in human_occurrences if item.meta_category
    )
    cluster_counts = Counter(cluster for item in occurrences for cluster in item.clusters)

    year_groups: dict[str, list[ModuleProfile]] = defaultdict(list)
    semester_groups: dict[str, list[ModuleProfile]] = defaultdict(list)
    stream_groups: dict[str, list[ModuleProfile]] = defaultdict(list)
    for module in modules:
        year_groups[module.year].append(module)
        if module.semester:
            semester_groups[module.semester].append(module)
        stream_groups[module.stream].append(module)

    def sort_group(group: dict[str, list[ModuleProfile]]) -> dict[str, list[ModuleProfile]]:
        return {
            key: sorted(value, key=lambda m: natural_sort_key(m.code))
            for key, value in sorted(group.items(), key=lambda kv: natural_sort_key(kv[0]))
        }

    return Analysis(
        modules=sorted(modules, key=lambda m: (natural_sort_key(m.year), natural_sort_key(m.semester), natural_sort_key(m.stream), natural_sort_key(m.code))),
        occurrences=occurrences,
        unique_skills=unique_skills,
        domain_occurrences=domain_occurrences,
        domain_unique_skills=domain_unique_skills,
        domain_modules=domain_modules,
        domain_years=domain_years,
        domain_streams=domain_streams,
        bloom_occurrences=bloom_occurrences,
        ai_occurrences=ai_occurrences,
        ai_unique_skills=ai_unique_skills,
        ai_domain_counts=ai_domain_counts,
        human_occurrences=human_occurrences,
        human_unique_skills=human_unique_skills,
        human_tag_counts=human_tag_counts,
        meta_category_counts=meta_category_counts,
        cluster_counts=cluster_counts,
        year_groups=sort_group(year_groups),
        semester_groups=sort_group(semester_groups),
        stream_groups=sort_group(stream_groups),
    )


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------
def chart_path(assets_dir: Path, name: str) -> Path:
    assets_dir.mkdir(parents=True, exist_ok=True)
    return assets_dir / name


def finish_chart(path: Path) -> Path:
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


def create_domain_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    rows = []
    for code, skills in analysis.domain_unique_skills.items():
        sample = next((x for x in analysis.occurrences if x.domain_code == code), None)
        label = f"{code} · {sample.domain_name if sample else code}"
        rows.append((label, len(skills)))
    rows.sort(key=lambda item: (item[1], item[0]))
    if not rows:
        return None
    labels = [textwrap.shorten(label, width=38, placeholder="…") for label, _ in rows]
    values = [value for _, value in rows]
    plt.figure(figsize=(9.2, max(4.2, 0.34 * len(rows))))
    plt.barh(labels, values)
    plt.xlabel(tr(language, "unique_skills"))
    plt.title(tr(language, "domain_inventory"))
    return finish_chart(chart_path(assets_dir, "domain_coverage.png"))


def create_bloom_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    keys = [key for key in BLOOM_ORDER if analysis.bloom_occurrences.get(key)]
    if analysis.bloom_occurrences.get("unknown"):
        keys.append("unknown")
    if not keys:
        return None
    values = [analysis.bloom_occurrences[key] for key in keys]
    labels = [bloom_label(key, language) for key in keys]
    plt.figure(figsize=(8.8, 4.8))
    plt.bar(labels, values)
    plt.ylabel(tr(language, "occurrences"))
    plt.title(tr(language, "bloom_profile"))
    plt.xticks(rotation=22, ha="right")
    return finish_chart(chart_path(assets_dir, "bloom_distribution.png"))


def create_ai_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    rows = sorted(analysis.ai_domain_counts.items(), key=lambda item: (item[1], item[0]))
    if not rows:
        return None
    labels = [label for label, _ in rows]
    values = [value for _, value in rows]
    plt.figure(figsize=(8.8, max(4.0, 0.45 * len(rows))))
    plt.barh(labels, values)
    plt.xlabel(tr(language, "occurrences"))
    plt.title(tr(language, "ai_profile"))
    return finish_chart(chart_path(assets_dir, "ai_domains.png"))


def create_human_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    combined = Counter(analysis.human_tag_counts)
    for category, count in analysis.meta_category_counts.items():
        combined[f"meta: {category}"] += count
    rows = sorted(combined.items(), key=lambda item: (item[1], item[0]))
    if not rows:
        return None
    labels = [label for label, _ in rows]
    values = [value for _, value in rows]
    plt.figure(figsize=(8.8, max(4.0, 0.45 * len(rows))))
    plt.barh(labels, values)
    plt.xlabel(tr(language, "occurrences"))
    plt.title(tr(language, "human_profile"))
    return finish_chart(chart_path(assets_dir, "human_capabilities.png"))


def create_bloom_progression_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    years = [year for year in analysis.year_groups if year and year != "Unspecified"]
    if len(years) < 1:
        return None
    years = sorted(years, key=natural_sort_key)
    series: dict[str, list[int]] = {key: [] for key in BLOOM_ORDER}
    for year in years:
        counter = Counter(
            outcome.bloom
            for module in analysis.year_groups[year]
            for outcome in module.outcomes
        )
        for key in BLOOM_ORDER:
            series[key].append(counter.get(key, 0))
    plt.figure(figsize=(9.2, 5.0))
    bottom = [0] * len(years)
    for key in BLOOM_ORDER:
        values = series[key]
        if not any(values):
            continue
        plt.bar(years, values, bottom=bottom, label=bloom_label(key, language))
        bottom = [a + b for a, b in zip(bottom, values)]
    plt.ylabel(tr(language, "occurrences"))
    plt.xlabel(tr(language, "year"))
    plt.title(tr(language, "by_year"))
    plt.legend(loc="upper left", bbox_to_anchor=(1.01, 1.0))
    return finish_chart(chart_path(assets_dir, "bloom_progression_by_year.png"))


def create_stream_matrix_chart(analysis: Analysis, assets_dir: Path, language: str) -> Path | None:
    domains = sorted(analysis.domain_unique_skills, key=natural_sort_key)
    streams = sorted(analysis.stream_groups, key=natural_sort_key)
    if not domains or not streams:
        return None
    matrix: list[list[int]] = []
    for domain in domains:
        row = []
        for stream in streams:
            skills = {
                item.skill_code
                for module in analysis.stream_groups[stream]
                for item in module.outcomes
                if item.domain_code == domain
            }
            row.append(len(skills))
        matrix.append(row)
    plt.figure(figsize=(max(8.2, 1.25 * len(streams)), max(5.0, 0.38 * len(domains))))
    image = plt.imshow(matrix, aspect="auto")
    plt.colorbar(image, label=tr(language, "unique_skills"))
    plt.xticks(range(len(streams)), [textwrap.shorten(x, 24, placeholder="…") for x in streams], rotation=35, ha="right")
    plt.yticks(range(len(domains)), domains)
    plt.title(tr(language, "coverage_matrix"))
    plt.xlabel(tr(language, "stream"))
    plt.ylabel(tr(language, "domain"))
    for i, row in enumerate(matrix):
        for j, value in enumerate(row):
            if value:
                plt.text(j, i, str(value), ha="center", va="center")
    return finish_chart(chart_path(assets_dir, "domain_stream_matrix.png"))


def create_charts(analysis: Analysis, assets_dir: Path, language: str) -> dict[str, Path]:
    chart_builders = {
        "domains": create_domain_chart,
        "bloom": create_bloom_chart,
        "ai": create_ai_chart,
        "human": create_human_chart,
        "bloom_year": create_bloom_progression_chart,
        "matrix": create_stream_matrix_chart,
    }
    charts: dict[str, Path] = {}
    for key, builder in chart_builders.items():
        path = builder(analysis, assets_dir, language)
        if path:
            charts[key] = path
    return charts


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, widths: Sequence[Cm]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    while len(grid):
        grid.remove(grid[0])
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width.cm * 567)))
        grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width.cm * 567)))
                tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "D7E0E2", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, bold: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(r_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def safe_bookmark(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_]", "_", value)
    if not value or value[0].isdigit():
        value = "m_" + value
    return value[:40]


def configure_document(document: Document, title: str, institution: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.06

    for style_name, size, color in (
        ("Title", 29, DARK_TEAL),
        ("Subtitle", 14, MID_GREY),
        ("Heading 1", 21, DARK_TEAL),
        ("Heading 2", 15, TEAL),
        ("Heading 3", 11.5, DARK_TEAL),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = institution
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        set_run_font(header.runs[0], size=8.5, bold=True, color=TEAL)
    footer = section.footer.paragraphs[0]
    footer.add_run(title + "   ·   ")
    set_run_font(footer.runs[0], size=8, color=MID_GREY)
    add_page_number(footer)

    document.core_properties.title = title
    document.core_properties.subject = "Curriculum skill-profile analysis"
    document.core_properties.author = institution
    document.core_properties.comments = "Generated from OOAPI module JSON and competence codebook data."
    set_update_fields(document)


def add_heading_bar(document: Document, text: str, level: int = 2):
    paragraph = document.add_heading(text, level=level)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), TEAL)
    borders.append(bottom)
    p_pr.append(borders)
    return paragraph


def add_callout(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_width(table, [Cm(15.5)])
    set_table_borders(table, color="B7CFD1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=DARK_TEAL)
    p = cell.add_paragraph(text)
    for run in p.runs:
        set_run_font(run, size=9.3, color=DARK_GREY)
    document.add_paragraph()


def add_chart(document: Document, path: Path | None, width_cm: float = 15.2) -> None:
    if not path or not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))


def add_cover(document: Document, analysis: Analysis, title: str, institution: str, programme: str, language: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(56)
    run = p.add_run(institution)
    set_run_font(run, size=13, bold=True, color=TEAL)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(title)
    set_run_font(run, size=30, bold=True, color=DARK_TEAL)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tr(language, "subtitle"))
    set_run_font(run, size=14, color=MID_GREY, italic=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(programme)
    set_run_font(run, size=12, bold=True, color=DARK_GREY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    run = p.add_run(f"{tr(language, 'assembled')}: {date.today().isoformat()}")
    set_run_font(run, size=10, color=MID_GREY)

    document.add_paragraph()
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(2.55)] * 6
    set_fixed_table_width(table, widths)
    values = [
        (tr(language, "modules"), len(analysis.modules)),
        (tr(language, "total_ects"), pretty_number(sum(m.ects for m in analysis.modules))),
        (tr(language, "unique_skills"), len(analysis.unique_skills)),
        (tr(language, "domains"), len(analysis.domain_unique_skills)),
        (tr(language, "ai_skills"), len(analysis.ai_unique_skills)),
        (tr(language, "human_skills"), len(analysis.human_unique_skills)),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, PALE_TEAL)
        set_cell_margins(cell, top=150, start=75, bottom=150, end=75)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        set_run_font(run, size=17, bold=True, color=DARK_TEAL)
        p = cell.add_paragraph(label)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=7.5, color=MID_GREY)
    set_table_borders(table, color="C9D8DA")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run(tr(language, "source_note"))
    set_run_font(run, size=9, italic=True, color=MID_GREY)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_contents(document: Document, language: str, has_semesters: bool) -> None:
    add_heading_bar(document, tr(language, "contents"), 1)
    sections = [
        tr(language, "module_index"),
        tr(language, "executive_profile"),
        tr(language, "domain_inventory"),
        tr(language, "bloom_profile"),
        tr(language, "ai_profile"),
        tr(language, "human_profile"),
        tr(language, "by_year"),
    ]
    if has_semesters:
        sections.append(tr(language, "by_semester"))
    sections.extend([
        tr(language, "by_stream"),
        tr(language, "coverage_matrix"),
        tr(language, "module_profiles"),
    ])
    for idx, section in enumerate(sections, start=1):
        p = document.add_paragraph(style="List Number")
        run = p.add_run(section)
        set_run_font(run, size=10, color=DARK_GREY)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def make_table(document: Document, headers: Sequence[str], widths: Sequence[Cm]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_width(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, TEAL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, size=8, bold=True, color=WHITE)
    return table


def add_table_row(table, values: Sequence[Any], *, shade: str | None = None, font_size: float = 8.2, bold_index: int | None = None) -> None:
    row = table.add_row()
    for idx, (cell, value) in enumerate(zip(row.cells, values)):
        if shade:
            set_cell_shading(cell, shade)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(clean_text(value))
        set_run_font(run, size=font_size, bold=(bold_index == idx), color=DARK_GREY)


def add_module_index(document: Document, analysis: Analysis, language: str) -> None:
    add_heading_bar(document, tr(language, "module_index"), 1)
    p = document.add_paragraph(tr(language, "source_note"))
    for run in p.runs:
        set_run_font(run, size=9, italic=True, color=MID_GREY)
    table = make_table(
        document,
        [tr(language, "code"), tr(language, "module"), tr(language, "stream"), tr(language, "year"), tr(language, "semester_short"), tr(language, "ects"), tr(language, "skills")],
        [Cm(1.7), Cm(4.2), Cm(3.4), Cm(1.1), Cm(1.2), Cm(1.2), Cm(2.7)],
    )
    for idx, module in enumerate(analysis.modules):
        row = table.add_row()
        values = [module.code, module.title, module.stream, module.year, module.semester or "—", pretty_number(module.ects), str(len(module.unique_skill_codes))]
        for cell_idx, (cell, value) in enumerate(zip(row.cells, values)):
            if idx % 2:
                set_cell_shading(cell, PALE_GREY)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            if cell_idx == 1:
                add_internal_hyperlink(p, value, safe_bookmark(module.code), bold=True)
            else:
                run = p.add_run(value)
                set_run_font(run, size=7.7, bold=(cell_idx == 0), color=DARK_GREY)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_executive_profile(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "executive_profile"), 1)
    add_callout(document, tr(language, "interpretation"), tr(language, "interpretation_text"))
    add_chart(document, charts.get("domains"))
    add_chart(document, charts.get("bloom"))
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_domain_inventory(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "domain_inventory"), 1)
    table = make_table(
        document,
        [tr(language, "domain"), tr(language, "domain_name"), tr(language, "skills"), tr(language, "modules_touched"), tr(language, "occurrences"), tr(language, "years"), tr(language, "streams")],
        [Cm(1.2), Cm(3.2), Cm(1.1), Cm(1.2), Cm(1.4), Cm(2.0), Cm(5.4)],
    )
    for idx, code in enumerate(sorted(analysis.domain_unique_skills, key=natural_sort_key)):
        sample = next((x for x in analysis.occurrences if x.domain_code == code), None)
        add_table_row(
            table,
            [
                code,
                sample.domain_name if sample else code,
                len(analysis.domain_unique_skills[code]),
                len(analysis.domain_modules[code]),
                analysis.domain_occurrences[code],
                ", ".join(unique_sorted(analysis.domain_years[code])),
                ", ".join(unique_sorted(analysis.domain_streams[code])),
            ],
            shade=PALE_GREY if idx % 2 else None,
            font_size=7.6,
            bold_index=0,
        )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_bloom_profile(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "bloom_profile"), 1)
    add_chart(document, charts.get("bloom"))
    add_chart(document, charts.get("bloom_year"))

    table = make_table(
        document,
        [tr(language, "bloom"), tr(language, "occurrences"), tr(language, "skills"), tr(language, "modules_touched")],
        [Cm(6.2), Cm(2.4), Cm(2.4), Cm(4.5)],
    )
    bloom_keys = [key for key in BLOOM_ORDER + ["unknown"] if analysis.bloom_occurrences.get(key)]
    for idx, key in enumerate(bloom_keys):
        unique = {item.skill_code for item in analysis.occurrences if item.bloom == key}
        modules = {item.module_code for item in analysis.occurrences if item.bloom == key}
        add_table_row(
            table,
            [bloom_label(key, language), analysis.bloom_occurrences[key], len(unique), len(modules)],
            shade=PALE_GREY if idx % 2 else None,
        )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_ai_profile(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "ai_profile"), 1)
    add_callout(document, tr(language, "interpretation"), tr(language, "ai_interpretation"))
    if not analysis.ai_unique_skills:
        document.add_paragraph(tr(language, "no_data"))
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return
    add_chart(document, charts.get("ai"))
    table = make_table(
        document,
        [tr(language, "code"), tr(language, "skill"), tr(language, "domain"), tr(language, "bloom"), tr(language, "ai_domains"), tr(language, "modules_touched")],
        [Cm(1.8), Cm(4.5), Cm(1.2), Cm(2.5), Cm(2.7), Cm(2.8)],
    )
    for idx, (code, sample) in enumerate(sorted(analysis.ai_unique_skills.items(), key=lambda kv: natural_sort_key(kv[0]))):
        modules = unique_sorted(item.module_code for item in analysis.ai_occurrences if item.skill_code == code)
        ai_domains = unique_sorted(domain for item in analysis.ai_occurrences if item.skill_code == code for domain in item.ai_domains)
        add_table_row(
            table,
            [code, sample.skill_name, sample.domain_code, bloom_label(sample.bloom, language), ", ".join(ai_domains) or "—", ", ".join(modules)],
            shade=PALE_GREY if idx % 2 else None,
            font_size=7.6,
            bold_index=0,
        )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_human_profile(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "human_profile"), 1)
    add_callout(document, tr(language, "interpretation"), tr(language, "human_interpretation"))
    if not analysis.human_unique_skills:
        document.add_paragraph(tr(language, "no_data"))
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return
    add_chart(document, charts.get("human"))
    table = make_table(
        document,
        [tr(language, "code"), tr(language, "skill"), tr(language, "human_tag"), tr(language, "meta_category"), tr(language, "clusters"), tr(language, "modules_touched")],
        [Cm(1.8), Cm(4.0), Cm(2.1), Cm(2.2), Cm(2.5), Cm(2.9)],
    )
    for idx, (code, sample) in enumerate(sorted(analysis.human_unique_skills.items(), key=lambda kv: natural_sort_key(kv[0]))):
        matching = [item for item in analysis.human_occurrences if item.skill_code == code]
        modules = unique_sorted(item.module_code for item in matching)
        tags = unique_sorted(item.human_tag for item in matching)
        categories = unique_sorted(item.meta_category for item in matching)
        clusters = unique_sorted(cluster for item in matching for cluster in item.clusters)
        add_table_row(
            table,
            [code, sample.skill_name, ", ".join(tags) or "—", ", ".join(categories) or "—", ", ".join(clusters) or "—", ", ".join(modules)],
            shade=PALE_GREY if idx % 2 else None,
            font_size=7.4,
            bold_index=0,
        )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def group_summary(modules: Sequence[ModuleProfile]) -> tuple[int, float, int, int, int, int]:
    occurrences = [item for module in modules for item in module.outcomes]
    unique_skills = {item.skill_code for item in occurrences}
    domains = {item.domain_code for item in occurrences}
    ai = {item.skill_code for item in occurrences if item.ai_related}
    human = {item.skill_code for item in occurrences if item.human_tag or item.meta_skill or item.meta_category}
    return len(modules), sum(m.ects for m in modules), len(unique_skills), len(domains), len(ai), len(human)


def add_group_section(document: Document, title: str, groups: Mapping[str, Sequence[ModuleProfile]], language: str, page_break: bool = True) -> None:
    add_heading_bar(document, title, 1)
    table = make_table(
        document,
        [title.split()[-1] if title else tr(language, "coverage"), tr(language, "modules"), tr(language, "total_ects"), tr(language, "unique_skills"), tr(language, "domains"), tr(language, "ai_skills"), tr(language, "human_skills")],
        [Cm(3.6), Cm(1.6), Cm(1.8), Cm(2.1), Cm(1.8), Cm(2.2), Cm(2.4)],
    )
    for idx, (key, modules) in enumerate(groups.items()):
        summary = group_summary(modules)
        add_table_row(
            table,
            [key, summary[0], pretty_number(summary[1]), summary[2], summary[3], summary[4], summary[5]],
            shade=PALE_GREY if idx % 2 else None,
            bold_index=0,
        )

    for key, modules in groups.items():
        document.add_heading(key, level=2)
        occurrences = [item for module in modules for item in module.outcomes]
        domain_counts = Counter(item.domain_code for item in occurrences)
        bloom_counts = Counter(item.bloom for item in occurrences)
        ai_domains = Counter(domain for item in occurrences if item.ai_related for domain in (item.ai_domains or ("unspecified",)))
        p = document.add_paragraph()
        p.add_run(f"{tr(language, 'modules')}: ").bold = True
        p.add_run(", ".join(module.code for module in modules))
        p = document.add_paragraph()
        p.add_run(f"{tr(language, 'domains')}: ").bold = True
        p.add_run(", ".join(f"{code} ({count})" for code, count in domain_counts.most_common()))
        p = document.add_paragraph()
        p.add_run(f"{tr(language, 'bloom')}: ").bold = True
        p.add_run(", ".join(f"{bloom_label(code, language)} ({count})" for code, count in bloom_counts.items()))
        if ai_domains:
            p = document.add_paragraph()
            p.add_run(f"{tr(language, 'ai_domains')}: ").bold = True
            p.add_run(", ".join(f"{name} ({count})" for name, count in ai_domains.most_common()))
    if page_break:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_coverage_matrix(document: Document, analysis: Analysis, charts: Mapping[str, Path], language: str) -> None:
    add_heading_bar(document, tr(language, "coverage_matrix"), 1)
    add_chart(document, charts.get("matrix"))
    streams = sorted(analysis.stream_groups, key=natural_sort_key)
    domains = sorted(analysis.domain_unique_skills, key=natural_sort_key)
    if streams and domains and len(streams) <= 7:
        widths = [Cm(2.2)] + [Cm(13.3 / len(streams)) for _ in streams]
        table = make_table(document, [tr(language, "domain")] + streams, widths)
        for idx, domain in enumerate(domains):
            values: list[Any] = [domain]
            for stream in streams:
                skills = {
                    item.skill_code
                    for module in analysis.stream_groups[stream]
                    for item in module.outcomes
                    if item.domain_code == domain
                }
                values.append(len(skills) or "—")
            add_table_row(table, values, shade=PALE_GREY if idx % 2 else None, font_size=7.4, bold_index=0)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_module_profiles(document: Document, analysis: Analysis, language: str) -> None:
    add_heading_bar(document, tr(language, "module_profiles"), 1)
    bookmark_id = 1000
    for index, module in enumerate(analysis.modules):
        heading = document.add_heading(f"{module.code} · {module.title}", level=2)
        add_bookmark(heading, safe_bookmark(module.code), bookmark_id)
        bookmark_id += 1

        # Profile strip
        table = document.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_fixed_table_width(table, [Cm(2.6)] * 6)
        labels_values = [
            (tr(language, "stream"), module.stream),
            (tr(language, "year"), module.year),
            (tr(language, "semester"), module.semester or "—"),
            (tr(language, "ects"), pretty_number(module.ects)),
            (tr(language, "skills"), len(module.unique_skill_codes)),
            (tr(language, "ai_skills"), len(module.ai_skills)),
        ]
        for cell, (label, value) in zip(table.rows[0].cells, labels_values):
            set_cell_shading(cell, PALE_TEAL)
            set_cell_margins(cell, top=100, start=70, bottom=100, end=70)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=10, bold=True, color=DARK_TEAL)
            p = cell.add_paragraph(label)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=6.8, color=MID_GREY)
        set_table_borders(table, color="C9D8DA")

        if module.description:
            p = document.add_paragraph(module.description)
            p.paragraph_format.space_before = Pt(5)

        domain_counts = Counter(item.domain_code for item in module.outcomes)
        bloom_counts = Counter(item.bloom for item in module.outcomes)
        p = document.add_paragraph()
        run = p.add_run(f"{tr(language, 'domains')}: ")
        set_run_font(run, size=9, bold=True, color=DARK_TEAL)
        p.add_run(", ".join(f"{code} ({count})" for code, count in domain_counts.most_common()) or "—")
        p = document.add_paragraph()
        run = p.add_run(f"{tr(language, 'bloom')}: ")
        set_run_font(run, size=9, bold=True, color=DARK_TEAL)
        p.add_run(", ".join(f"{bloom_label(key, language)} ({count})" for key, count in bloom_counts.items()) or "—")

        if module.ai_skills:
            document.add_heading(tr(language, "ai_profile"), level=3)
            for item in sorted((x for x in module.outcomes if x.ai_related), key=lambda x: natural_sort_key(x.skill_code)):
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(f"{item.skill_code} · {item.skill_name}")
                set_run_font(run, size=8.8, bold=True, color=DARK_TEAL)
                if item.ai_domains:
                    p.add_run(f" — {', '.join(item.ai_domains)}")

        human_items = [x for x in module.outcomes if x.human_tag or x.meta_skill or x.meta_category]
        if human_items:
            document.add_heading(tr(language, "human_profile"), level=3)
            for item in sorted(human_items, key=lambda x: natural_sort_key(x.skill_code)):
                tags = unique_sorted([item.human_tag, item.meta_category, *item.clusters])
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(f"{item.skill_code} · {item.skill_name}")
                set_run_font(run, size=8.8, bold=True, color=DARK_TEAL)
                if tags:
                    p.add_run(f" — {', '.join(tags)}")

        document.add_heading(tr(language, "skills"), level=3)
        skill_table = make_table(
            document,
            [tr(language, "code"), tr(language, "skill"), tr(language, "domain"), tr(language, "bloom"), tr(language, "ai"), tr(language, "human")],
            [Cm(1.8), Cm(5.8), Cm(1.4), Cm(2.5), Cm(1.7), Cm(2.3)],
        )
        for row_idx, item in enumerate(sorted(module.outcomes, key=lambda x: (natural_sort_key(x.domain_code), natural_sort_key(x.skill_code)))):
            human_label = ", ".join(unique_sorted([item.human_tag, item.meta_category])) or "—"
            add_table_row(
                skill_table,
                [item.skill_code, item.skill_name, item.domain_code, bloom_label(item.bloom, language), "✓" if item.ai_related else "—", human_label],
                shade=PALE_GREY if row_idx % 2 else None,
                font_size=7.4,
                bold_index=0,
            )
        p = document.add_paragraph()
        add_internal_hyperlink(p, tr(language, "back"), "module_index")
        if index < len(analysis.modules) - 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_docx(
    analysis: Analysis,
    charts: Mapping[str, Path],
    output: Path,
    title: str,
    institution: str,
    programme: str,
    language: str,
) -> None:
    document = Document()
    configure_document(document, title, institution)
    add_cover(document, analysis, title, institution, programme, language)
    add_contents(document, language, bool(analysis.semester_groups))

    # Bookmark on index heading for back-links.
    index_heading = add_heading_bar(document, tr(language, "module_index"), 1)
    add_bookmark(index_heading, "module_index", 1)
    # Recreate index table without adding a second heading.
    table = make_table(
        document,
        [tr(language, "code"), tr(language, "module"), tr(language, "stream"), tr(language, "year"), tr(language, "semester_short"), tr(language, "ects"), tr(language, "skills")],
        [Cm(1.7), Cm(4.2), Cm(3.4), Cm(1.1), Cm(1.2), Cm(1.2), Cm(2.7)],
    )
    for idx, module in enumerate(analysis.modules):
        row = table.add_row()
        values = [module.code, module.title, module.stream, module.year, module.semester or "—", pretty_number(module.ects), str(len(module.unique_skill_codes))]
        for cell_idx, (cell, value) in enumerate(zip(row.cells, values)):
            if idx % 2:
                set_cell_shading(cell, PALE_GREY)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            if cell_idx == 1:
                add_internal_hyperlink(p, value, safe_bookmark(module.code), bold=True)
            else:
                run = p.add_run(value)
                set_run_font(run, size=7.7, bold=(cell_idx == 0), color=DARK_GREY)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_executive_profile(document, analysis, charts, language)
    add_domain_inventory(document, analysis, charts, language)
    add_bloom_profile(document, analysis, charts, language)
    add_ai_profile(document, analysis, charts, language)
    add_human_profile(document, analysis, charts, language)
    add_group_section(document, tr(language, "by_year"), analysis.year_groups, language)
    if analysis.semester_groups:
        add_group_section(document, tr(language, "by_semester"), analysis.semester_groups, language)
    add_group_section(document, tr(language, "by_stream"), analysis.stream_groups, language)
    add_coverage_matrix(document, analysis, charts, language)
    add_module_profiles(document, analysis, language)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


# ---------------------------------------------------------------------------
# Markdown output
# ---------------------------------------------------------------------------
def md_table(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> list[str]:
    lines = [
        "| " + " | ".join(markdown_escape(x) for x in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    lines.extend("| " + " | ".join(markdown_escape(x) for x in row) + " |" for row in rows)
    return lines


def chart_md(path: Path | None, markdown_output: Path) -> str:
    if not path:
        return ""
    try:
        rel = path.relative_to(markdown_output.parent)
    except ValueError:
        rel = path
    return f"![{path.stem}]({rel.as_posix()})"


def build_markdown(
    analysis: Analysis,
    charts: Mapping[str, Path],
    output: Path,
    title: str,
    institution: str,
    programme: str,
    language: str,
) -> None:
    lines: list[str] = [
        "---",
        'layout: page',
        f'title: "{title}"',
        "---",
        "",
        f"# {title}",
        "",
        f"**{institution}**  ",
        f"{programme}  ",
        f"{tr(language, 'assembled')}: {date.today().isoformat()}",
        "",
        f"> {tr(language, 'source_note')}",
        "",
        f"## {tr(language, 'portfolio_snapshot')}",
        "",
    ]
    snapshot = [
        (tr(language, "modules"), len(analysis.modules)),
        (tr(language, "total_ects"), pretty_number(sum(m.ects for m in analysis.modules))),
        (tr(language, "unique_skills"), len(analysis.unique_skills)),
        (tr(language, "domains"), len(analysis.domain_unique_skills)),
        (tr(language, "ai_skills"), len(analysis.ai_unique_skills)),
        (tr(language, "human_skills"), len(analysis.human_unique_skills)),
    ]
    lines += md_table(["Measure", "Value"], snapshot)

    lines += ["", '<a id="module-index"></a>', "", f"## {tr(language, 'module_index')}", ""]
    module_rows = [
        [
            module.code,
            f"[{module.title}](#{markdown_anchor(module.code)})",
            module.stream,
            module.year,
            module.semester or "—",
            pretty_number(module.ects),
            len(module.unique_skill_codes),
        ]
        for module in analysis.modules
    ]
    # Preserve links by constructing the module table manually for the linked column.
    lines.append(f"| {tr(language, 'code')} | {tr(language, 'module')} | {tr(language, 'stream')} | {tr(language, 'year')} | {tr(language, 'semester')} | {tr(language, 'ects')} | {tr(language, 'skills')} |")
    lines.append("| --- | --- | --- | --- | --- | --- | --- |")
    for row in module_rows:
        lines.append("| " + " | ".join(markdown_escape(x) if idx != 1 else str(x) for idx, x in enumerate(row)) + " |")

    lines += [
        "",
        f"## {tr(language, 'executive_profile')}",
        "",
        f"> **{tr(language, 'interpretation')}** — {tr(language, 'interpretation_text')}",
        "",
    ]
    for key in ("domains", "bloom"):
        chart = chart_md(charts.get(key), output)
        if chart:
            lines += [chart, ""]

    lines += ["", f"## {tr(language, 'domain_inventory')}", ""]
    domain_rows = []
    for code in sorted(analysis.domain_unique_skills, key=natural_sort_key):
        sample = next((x for x in analysis.occurrences if x.domain_code == code), None)
        domain_rows.append([
            code,
            sample.domain_name if sample else code,
            len(analysis.domain_unique_skills[code]),
            len(analysis.domain_modules[code]),
            analysis.domain_occurrences[code],
            ", ".join(unique_sorted(analysis.domain_years[code])),
            ", ".join(unique_sorted(analysis.domain_streams[code])),
        ])
    lines += md_table(
        [tr(language, "domain"), tr(language, "domain_name"), tr(language, "skills"), tr(language, "modules_touched"), tr(language, "occurrences"), tr(language, "years"), tr(language, "streams")],
        domain_rows,
    )

    lines += ["", f"## {tr(language, 'bloom_profile')}", ""]
    for key in ("bloom", "bloom_year"):
        chart = chart_md(charts.get(key), output)
        if chart:
            lines += [chart, ""]
    bloom_rows = []
    for key in BLOOM_ORDER + ["unknown"]:
        if not analysis.bloom_occurrences.get(key):
            continue
        bloom_rows.append([
            bloom_label(key, language),
            analysis.bloom_occurrences[key],
            len({x.skill_code for x in analysis.occurrences if x.bloom == key}),
            len({x.module_code for x in analysis.occurrences if x.bloom == key}),
        ])
    lines += md_table([tr(language, "bloom"), tr(language, "occurrences"), tr(language, "skills"), tr(language, "modules_touched")], bloom_rows)

    lines += ["", f"## {tr(language, 'ai_profile')}", "", f"> {tr(language, 'ai_interpretation')}", ""]
    ai_chart = chart_md(charts.get("ai"), output)
    if ai_chart:
        lines += [ai_chart, ""]
    ai_rows = []
    for code, sample in sorted(analysis.ai_unique_skills.items(), key=lambda kv: natural_sort_key(kv[0])):
        matching = [x for x in analysis.ai_occurrences if x.skill_code == code]
        ai_rows.append([
            code,
            sample.skill_name,
            sample.domain_code,
            bloom_label(sample.bloom, language),
            ", ".join(unique_sorted(d for x in matching for d in x.ai_domains)) or "—",
            ", ".join(unique_sorted(x.module_code for x in matching)),
        ])
    lines += md_table([tr(language, "code"), tr(language, "skill"), tr(language, "domain"), tr(language, "bloom"), tr(language, "ai_domains"), tr(language, "modules_touched")], ai_rows) if ai_rows else [tr(language, "no_data")]

    lines += ["", f"## {tr(language, 'human_profile')}", "", f"> {tr(language, 'human_interpretation')}", ""]
    human_chart = chart_md(charts.get("human"), output)
    if human_chart:
        lines += [human_chart, ""]
    human_rows = []
    for code, sample in sorted(analysis.human_unique_skills.items(), key=lambda kv: natural_sort_key(kv[0])):
        matching = [x for x in analysis.human_occurrences if x.skill_code == code]
        human_rows.append([
            code,
            sample.skill_name,
            ", ".join(unique_sorted(x.human_tag for x in matching)) or "—",
            ", ".join(unique_sorted(x.meta_category for x in matching)) or "—",
            ", ".join(unique_sorted(c for x in matching for c in x.clusters)) or "—",
            ", ".join(unique_sorted(x.module_code for x in matching)),
        ])
    lines += md_table([tr(language, "code"), tr(language, "skill"), tr(language, "human_tag"), tr(language, "meta_category"), tr(language, "clusters"), tr(language, "modules_touched")], human_rows) if human_rows else [tr(language, "no_data")]

    def add_group_md(title_text: str, groups: Mapping[str, Sequence[ModuleProfile]]) -> None:
        lines.extend(["", f"## {title_text}", ""])
        rows = []
        for key, modules in groups.items():
            summary = group_summary(modules)
            rows.append([key, summary[0], pretty_number(summary[1]), summary[2], summary[3], summary[4], summary[5]])
        lines.extend(md_table([title_text, tr(language, "modules"), tr(language, "total_ects"), tr(language, "unique_skills"), tr(language, "domains"), tr(language, "ai_skills"), tr(language, "human_skills")], rows))
        for key, modules in groups.items():
            occurrences = [item for module in modules for item in module.outcomes]
            domains = Counter(item.domain_code for item in occurrences)
            blooms = Counter(item.bloom for item in occurrences)
            lines.extend([
                "",
                f"### {key}",
                "",
                f"**{tr(language, 'modules')}:** {', '.join(m.code for m in modules)}  ",
                f"**{tr(language, 'domains')}:** {', '.join(f'{k} ({v})' for k, v in domains.most_common()) or '—'}  ",
                f"**{tr(language, 'bloom')}:** {', '.join(f'{bloom_label(k, language)} ({v})' for k, v in blooms.items()) or '—'}",
            ])

    add_group_md(tr(language, "by_year"), analysis.year_groups)
    if analysis.semester_groups:
        add_group_md(tr(language, "by_semester"), analysis.semester_groups)
    add_group_md(tr(language, "by_stream"), analysis.stream_groups)

    lines += ["", f"## {tr(language, 'coverage_matrix')}", ""]
    matrix_chart = chart_md(charts.get("matrix"), output)
    if matrix_chart:
        lines += [matrix_chart, ""]

    lines += ["", f"## {tr(language, 'module_profiles')}", ""]
    for module in analysis.modules:
        lines += [
            f'<a id="{markdown_anchor(module.code)}"></a>',
            "",
            f"### {module.code} · {module.title}",
            "",
            f"**{tr(language, 'stream')}:** {module.stream}  ",
            f"**{tr(language, 'year')}:** {module.year}  ",
            f"**{tr(language, 'semester')}:** {module.semester or '—'}  ",
            f"**{tr(language, 'ects')}:** {pretty_number(module.ects)}  ",
            f"**{tr(language, 'source')}:** `{module.path.name}`",
            "",
        ]
        if module.description:
            lines += [module.description, ""]
        rows = [
            [
                item.skill_code,
                item.skill_name,
                item.domain_code,
                bloom_label(item.bloom, language),
                "Yes" if item.ai_related else "—",
                ", ".join(unique_sorted([item.human_tag, item.meta_category])) or "—",
            ]
            for item in sorted(module.outcomes, key=lambda x: (natural_sort_key(x.domain_code), natural_sort_key(x.skill_code)))
        ]
        lines += md_table([tr(language, "code"), tr(language, "skill"), tr(language, "domain"), tr(language, "bloom"), tr(language, "ai"), tr(language, "human")], rows)
        lines += ["", f"[↑ {tr(language, 'back')}](#module-index)", ""]

    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create a Word and Markdown skill-profile atlas from OOAPI module JSON files."
    )
    parser.add_argument("--dir", type=Path, required=True, help="Directory containing module JSON files")
    parser.add_argument("--pattern", default="T*INF*v6.json", help="Glob pattern for selected module files")
    parser.add_argument("--codebook", type=Path, required=True, help="Competence codebook JSON")
    parser.add_argument("--output", type=Path, default=Path("OUG_Skill_Profile_Atlas.docx"), help="Word output path")
    parser.add_argument("--markdown-output", type=Path, help="Markdown output path; defaults to the Word output stem")
    parser.add_argument("--assets-dir", type=Path, help="Directory for generated charts; defaults beside the output")
    parser.add_argument("--language", default="en-GB", choices=["en-GB", "de-DE"], help="Preferred output language")
    parser.add_argument("--title", default="OUG Skill Profile Atlas", help="Document title")
    parser.add_argument("--institution", default="Open University of Germany (OUG)", help="Institution name")
    parser.add_argument("--programme", default="Bachelor of Computer Science (B.Sc.)", help="Programme name")
    parser.add_argument("--docx-only", action="store_true", help="Create only the Word report")
    parser.add_argument("--markdown-only", action="store_true", help="Create only the Markdown report")
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    if args.docx_only and args.markdown_only:
        print("Error: --docx-only and --markdown-only cannot be used together.", file=sys.stderr)
        return 2
    try:
        codebook_data = load_json(args.codebook)
        codebook = CodebookIndex(codebook_data, args.language)
        modules = load_modules(args.dir, args.pattern, codebook, args.language)
        analysis = analyse(modules)

        markdown_output = args.markdown_output or args.output.with_suffix(".md")
        assets_dir = args.assets_dir or args.output.parent / f"{args.output.stem}_assets"
        charts = create_charts(analysis, assets_dir, args.language)

        if not args.markdown_only:
            build_docx(
                analysis,
                charts,
                args.output,
                args.title,
                args.institution,
                args.programme,
                args.language,
            )
            print(f"Created Word report: {args.output.resolve()}")
        if not args.docx_only:
            build_markdown(
                analysis,
                charts,
                markdown_output,
                args.title,
                args.institution,
                args.programme,
                args.language,
            )
            print(f"Created Markdown report: {markdown_output.resolve()}")
        print(f"Created chart assets: {assets_dir.resolve()}")
        print(
            "Summary: "
            f"{len(analysis.modules)} modules, "
            f"{len(analysis.unique_skills)} unique skills, "
            f"{len(analysis.domain_unique_skills)} domains, "
            f"{len(analysis.ai_unique_skills)} AI-related skills, "
            f"{len(analysis.human_unique_skills)} human/meta-tagged skills."
        )
        return 0
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
